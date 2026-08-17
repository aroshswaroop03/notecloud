"""
app.py — Scrib-d Flask Backend
Receives an uploaded image, sends it to Claude via the Anthropic API,
and returns the transcribed handwriting as plain text.

Now also includes a full user authentication system using SQLite (a file-based
database built right into Python — no separate server needed) and Flask sessions
(a way to remember who is logged in across multiple page requests).
"""

import base64
import gzip
import io
import json
import os
import re
import secrets
import smtplib
import sqlite3
from datetime import date, datetime, timedelta
from email.message import EmailMessage
from functools import wraps
from urllib.parse import urlparse

import anthropic
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from flask_wtf.csrf import CSRFProtect
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
try:
    from PIL import Image as PilImage
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
try:
    from google_auth_oauthlib.flow import Flow
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request as GoogleRequest
    from googleapiclient.discovery import build as google_build
    GOOGLE_LIBS_AVAILABLE = True
except ImportError:
    GOOGLE_LIBS_AVAILABLE = False
try:
    import stripe
    STRIPE_LIBS_AVAILABLE = True
except ImportError:
    STRIPE_LIBS_AVAILABLE = False
try:
    import requests as http_requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False
try:
    import jwt as pyjwt
    APPLE_LIBS_AVAILABLE = True
except ImportError:
    APPLE_LIBS_AVAILABLE = False
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from flask import (
    Flask,
    Response,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from werkzeug.security import check_password_hash, generate_password_hash

# Load environment variables from the .env file so ANTHROPIC_API_KEY is available
load_dotenv()
os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"

# Create the Flask application. Flask looks for templates in a "templates/" folder
# and static files (CSS, JS, images) in a "static/" folder by default.
app = Flask(__name__)

# ── Secret key ────────────────────────────────────────────────────────────────
app.secret_key = os.getenv("SECRET_KEY") or secrets.token_hex(32)
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SECURE"] = os.getenv("FLASK_ENV") == "production"

# ── Security extensions ───────────────────────────────────────────────────────
csrf    = CSRFProtect(app)
limiter = Limiter(get_remote_address, app=app, default_limits=[])

# ── Singleton Anthropic client (avoids re-initialising the HTTP client per request) ──
_anthropic = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# Return JSON for CSRF errors (instead of HTML which the frontend can't parse)
from flask_wtf.csrf import CSRFError

@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    return jsonify({"error": "Request validation failed. Please refresh the page and try again."}), 400

@app.errorhandler(500)
def handle_500(e):
    return jsonify({"error": "Internal server error. Please try again."}), 500

@app.errorhandler(429)
def handle_429(e):
    return jsonify({"error": "Too many requests. Please wait a moment and try again."}), 429


# ── Response post-processing: gzip + caching ──────────────────────────────────
# Compressible text types worth gzipping. The big win is the ~200KB index.html,
# which shrinks to roughly 30KB on the wire.
_COMPRESSIBLE = ("text/html", "text/css", "application/javascript",
                 "application/json", "image/svg+xml", "text/plain")

@app.template_global()
def asset(filename):
    """
    URL for a file in /static with a cache-busting ?v=<mtime> appended.

    Static responses carry a 30-day max-age (see _optimize_response below), so
    without this a changed logo or favicon stays stale in browsers that already
    have it — they never revalidate. Keying on the file's mtime means editing
    the file is all it takes to invalidate it.
    """
    try:
        version = int(os.path.getmtime(os.path.join(app.static_folder, filename)))
    except OSError:
        return f"/static/{filename}"
    return f"/static/{filename}?v={version}"


@app.after_request
def _optimize_response(resp):
    # Baseline security headers (defense-in-depth):
    #   nosniff        — stop the browser MIME-sniffing a response into script
    #   X-Frame-Options— block the site being iframed elsewhere (clickjacking)
    #   Referrer-Policy— don't leak full URLs to third parties
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    resp.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")

    # Long-lived caching for static assets (avatars, logos) so the browser
    # doesn't re-fetch them on every page load. Avatar uploads already append a
    # ?t=<timestamp> cache-buster, so a long max-age is safe.
    if request.path.startswith("/static/") and resp.status_code == 200:
        resp.headers["Cache-Control"] = "public, max-age=2592000"  # 30 days

    # gzip text responses when the client supports it and the body is big enough
    # to be worth it (tiny bodies cost more in CPU than they save on the wire).
    try:
        accepts = request.headers.get("Accept-Encoding", "")
        ctype = (resp.content_type or "").split(";")[0].strip()
        if ("gzip" in accepts
                and resp.direct_passthrough is False
                and ctype in _COMPRESSIBLE
                and "Content-Encoding" not in resp.headers):
            data = resp.get_data()
            if len(data) >= 1024:
                compressed = gzip.compress(data, compresslevel=6)
                resp.set_data(compressed)
                resp.headers["Content-Encoding"] = "gzip"
                resp.headers["Vary"] = "Accept-Encoding"
                resp.headers["Content-Length"] = str(len(compressed))
    except Exception:
        # Never let compression break a response — fall back to uncompressed.
        pass
    return resp

# ── Database path ──────────────────────────────────────────────────────────────
# __file__ is the path to this script. os.path.dirname gets the folder it lives
# in. We store the database in the same folder as app.py.
DB_PATH = os.path.join(os.path.dirname(__file__), "scrib_d.db")

# ── Outgoing email ────────────────────────────────────────────────────────────
# Plain SMTP via the standard library, so any provider works (Gmail app
# password, SendGrid, Mailgun, SES, Postmark — they all speak SMTP).
# Leave these unset in dev: send_email() then prints the message to the console
# instead of sending it, which is enough to run the reset flow end to end.
EMAIL_HOST     = os.getenv("EMAIL_HOST", "")
EMAIL_PORT     = int(os.getenv("EMAIL_PORT", "587"))
EMAIL_USER     = os.getenv("EMAIL_USER", "")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD", "")
EMAIL_FROM     = os.getenv("EMAIL_FROM", "NoteCloud <no-reply@note-cloud.com>")

# ── Password reset ────────────────────────────────────────────────────────────
RESET_CODE_TTL_MINUTES = 10   # how long an emailed code stays valid
RESET_MAX_ATTEMPTS     = 5    # wrong guesses before a code is burned

# ── Login 2FA (email OTP, periodic) ───────────────────────────────────────────
# Only gates email+password login — Google/Apple sign-in already proves
# identity through an external provider, so it skips this entirely.
LOGIN_OTP_TTL_MINUTES  = 10
LOGIN_OTP_MAX_ATTEMPTS = 5
TRUSTED_DEVICE_DAYS    = 30    # how long a verified browser skips the code
TRUSTED_DEVICE_COOKIE  = "td"

# ── Token limits per tier ─────────────────────────────────────────────────────
# One "token" = one word in the transcription output.
# ~200-250 words per handwritten page, so:
#   Free    (500)  ≈ 2 pages/day
#   Student (5000) ≈ 20 pages/day
#   Pro     (None) = unlimited
TIER_LIMITS = {
    "free":    500,
    "student": 5000,
    "pro":     None,   # None = unlimited
}

# How many pages (images) a single upload may contain, per tier.
# None = unlimited. Admins and the "dev" tier are treated as unlimited.
TIER_PAGE_LIMITS = {
    "free":    1,
    "student": 5,
    "pro":     None,
}

# How many bonus tokens a referrer earns per person they bring in (daily)
REFERRAL_BONUS_TOKENS = 250  # roughly 1 extra page per referral

# The secret owner code — loaded from .env so it's never in the source code.
# Whoever redeems this code gets is_admin=1 and is never limited.
OWNER_CODE = os.getenv("OWNER_CODE", "")

# Google OAuth 2.0 credentials — set these in .env
GOOGLE_CLIENT_ID     = os.getenv("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET      = os.getenv("GOOGLE_CLIENT_SECRET", "")
GOOGLE_REDIRECT_URI       = os.getenv("GOOGLE_REDIRECT_URI", "http://127.0.0.1:5000/google/callback")
GOOGLE_LOGIN_REDIRECT_URI = os.getenv("GOOGLE_LOGIN_REDIRECT_URI", "http://127.0.0.1:5000/auth/google/callback")
GOOGLE_SCOPES = [
    "https://www.googleapis.com/auth/documents",
    "https://www.googleapis.com/auth/drive.file",
]
GOOGLE_LOGIN_SCOPES = [
    "openid",
    "https://www.googleapis.com/auth/userinfo.email",
    "https://www.googleapis.com/auth/userinfo.profile",
]

# Stripe — set these in .env once you have a Stripe account
STRIPE_SECRET_KEY     = os.getenv("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET", "")
STRIPE_PRICE_IDS = {
    ("student", "monthly"): os.getenv("STRIPE_PRICE_STUDENT_MONTHLY", ""),
    ("student", "annual"):  os.getenv("STRIPE_PRICE_STUDENT_ANNUAL", ""),
    ("pro", "monthly"):     os.getenv("STRIPE_PRICE_PRO_MONTHLY", ""),
    ("pro", "annual"):      os.getenv("STRIPE_PRICE_PRO_ANNUAL", ""),
}
# Reverse lookup so the webhook can turn "which price did they buy" back into a tier
STRIPE_PRICE_TO_TIER = {v: k[0] for k, v in STRIPE_PRICE_IDS.items() if v}
if STRIPE_LIBS_AVAILABLE and STRIPE_SECRET_KEY:
    stripe.api_key = STRIPE_SECRET_KEY


def send_email(to_address, subject, body):
    """
    Send a plain-text email. Returns True if it went out, False otherwise.

    With no EMAIL_HOST configured this logs the message instead of sending —
    that keeps local development working without credentials, and the caller
    treats both cases as success so behaviour doesn't diverge between dev and
    production. Never raises: a failure to send must not break the request.
    """
    if not EMAIL_HOST:
        app.logger.info(
            "email not configured — would have sent to %s:\nSubject: %s\n%s",
            to_address, subject, body,
        )
        return False

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"]    = EMAIL_FROM
    msg["To"]      = to_address
    msg.set_content(body)

    try:
        with smtplib.SMTP(EMAIL_HOST, EMAIL_PORT, timeout=15) as smtp:
            smtp.starttls()
            if EMAIL_USER:
                smtp.login(EMAIL_USER, EMAIL_PASSWORD)
            smtp.send_message(msg)
        return True
    except Exception as e:
        app.logger.error("email send failed to %s: %s", to_address, e)
        return False


def issue_login_otp(user):
    """
    Create a login 2FA code, email it, and return the plaintext code — the
    only place it exists unhashed. Mirrors the password-reset issuance, but
    into login_otps and with sign-in phrasing so the two flows read as
    distinct actions to anyone glancing at their inbox.
    """
    conn = get_db()
    code = f"{secrets.randbelow(1_000_000):06d}"
    now  = datetime.utcnow()

    conn.execute("UPDATE login_otps SET used = 1 WHERE user_id = ? AND used = 0", (user["id"],))
    conn.execute(
        """INSERT INTO login_otps (user_id, code_hash, expires_at, created_at)
           VALUES (?, ?, ?, ?)""",
        (
            user["id"],
            generate_password_hash(code, method="pbkdf2:sha256"),
            (now + timedelta(minutes=LOGIN_OTP_TTL_MINUTES)).isoformat(),
            now.isoformat(),
        ),
    )
    conn.commit()
    conn.close()

    name = user["first_name"] or "there"
    send_email(
        user["email"],
        "Your NoteCloud sign-in code",
        f"""Hi {name},

Someone is signing in to your NoteCloud account from a new browser. Your
sign-in code is:

    {code}

It expires in {LOGIN_OTP_TTL_MINUTES} minutes.

If this wasn't you, change your password — someone else has it.

— NoteCloud
""",
    )
    return code


def _trusted_device_serializer():
    # A distinct salt keeps this signature namespace-separate from Flask's own
    # session cookie and from any other itsdangerous use, so a token minted
    # for one purpose can never be replayed as another.
    return URLSafeTimedSerializer(app.secret_key, salt="trusted-device-v1")


def is_trusted_device(user_id):
    """True if this browser already completed 2FA for this exact account
    within the last TRUSTED_DEVICE_DAYS. A cookie minted for a different
    account (shared computer) does not count."""
    token = request.cookies.get(TRUSTED_DEVICE_COOKIE)
    if not token:
        return False
    try:
        seen_user_id = _trusted_device_serializer().loads(
            token, max_age=TRUSTED_DEVICE_DAYS * 86400
        )
    except (BadSignature, SignatureExpired):
        return False
    return seen_user_id == user_id


def mark_device_trusted(resp, user_id):
    """Set the cookie that lets this exact browser skip 2FA next time."""
    token = _trusted_device_serializer().dumps(user_id)
    resp.set_cookie(
        TRUSTED_DEVICE_COOKIE,
        token,
        max_age=TRUSTED_DEVICE_DAYS * 86400,
        httponly=True,
        samesite="Lax",
        secure=os.getenv("FLASK_ENV") == "production",
    )
    return resp


# Characters the URL spec strips before parsing (tab, CR, LF) — left in place,
# they let a value like "/\t/evil.com" slip past a naive same-site check and
# still resolve as "//evil.com" once a browser normalizes it away.
_URL_STRIP_CHARS = "\t\r\n"


def safe_next_path(next_url):
    """
    Only allow a same-site relative path for a post-OAuth redirect.

    A raw prefix check (startswith("/") and not startswith("//")) isn't
    enough on its own: browsers strip tab/CR/LF and normalize a leading
    backslash to a forward slash when resolving http(s) URLs (WHATWG URL
    spec), so "/\\evil.com" or "/\t/evil.com" pass that check as strings but
    still resolve to an attacker-controlled host. Reject anything containing
    those characters outright, and use urlparse to confirm there's no
    scheme/host hiding in what's left.
    """
    if not next_url:
        return "/"
    cleaned = "".join(ch for ch in next_url if ch not in _URL_STRIP_CHARS)
    if cleaned != next_url or "\\" in cleaned:
        return "/"
    parsed = urlparse(cleaned)
    if parsed.scheme or parsed.netloc:
        return "/"
    if not cleaned.startswith("/") or cleaned.startswith("//"):
        return "/"
    return cleaned


def tier_from_subscription(sub):
    """Map a subscription's current price back to one of our tier names."""
    try:
        return STRIPE_PRICE_TO_TIER.get(sub["items"]["data"][0]["price"]["id"])
    except (KeyError, IndexError, TypeError):
        return None


def subscription_period_end(sub):
    """
    Unix timestamp for when the current billing period ends, or None.

    Stripe moved current_period_end off the Subscription and onto each
    subscription item in the 2025-03-31 API version, so check both rather than
    assuming whichever version this account happens to be pinned to.
    """
    if not sub:
        return None
    end = sub.get("current_period_end")
    if end:
        return end
    try:
        return sub["items"]["data"][0].get("current_period_end")
    except (KeyError, IndexError, TypeError):
        return None

# Notion OAuth — set these in .env once you've created a public Notion integration
NOTION_CLIENT_ID     = os.getenv("NOTION_CLIENT_ID", "")
NOTION_CLIENT_SECRET = os.getenv("NOTION_CLIENT_SECRET", "")
NOTION_REDIRECT_URI  = os.getenv("NOTION_REDIRECT_URI", "http://127.0.0.1:5000/notion/callback")

# Sign in with Apple — set these in .env once you have an Apple Developer account
APPLE_CLIENT_ID    = os.getenv("APPLE_CLIENT_ID", "")   # the Services ID, e.g. com.notecloud.web
APPLE_TEAM_ID      = os.getenv("APPLE_TEAM_ID", "")
APPLE_KEY_ID       = os.getenv("APPLE_KEY_ID", "")
APPLE_PRIVATE_KEY  = os.getenv("APPLE_PRIVATE_KEY", "").replace("\\n", "\n")  # contents of the .p8 file
APPLE_REDIRECT_URI = os.getenv("APPLE_REDIRECT_URI", "http://127.0.0.1:5000/auth/apple/callback")

# The maximum upload size Flask will accept — 10 MB should be plenty for a photo
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB

# File types we're willing to accept. We reject anything else before it hits the AI.
ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp"}

# Map file extensions to the MIME types Claude expects
MIME_TYPES = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
}


# ── Database helpers ───────────────────────────────────────────────────────────

def get_db():
    """
    Open a connection to the SQLite database and return it.

    sqlite3.connect() creates the file automatically if it doesn't exist yet.
    row_factory = sqlite3.Row makes rows behave like dictionaries so we can
    access columns by name (e.g. row["email"]) instead of by index (row[1]).
    """
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    # Per-connection performance pragmas:
    #   synchronous=NORMAL — safe with WAL, far fewer fsyncs than FULL
    #   temp_store=MEMORY  — keep temp B-trees in RAM
    #   cache_size=-8000   — ~8MB page cache (negative = KB)
    #   busy_timeout       — wait up to 3s for a lock instead of erroring out
    conn.execute("PRAGMA synchronous = NORMAL")
    conn.execute("PRAGMA temp_store = MEMORY")
    conn.execute("PRAGMA cache_size = -8000")
    conn.execute("PRAGMA busy_timeout = 3000")
    return conn


def init_db():
    """
    Create the database tables if they don't already exist.

    This is called once when the app starts. The CREATE TABLE IF NOT EXISTS
    statement is safe to run every time — it only creates the table when it's
    missing, so existing data is never deleted.

    Table: users
      id              — auto-incrementing integer, the primary key
      email           — must be unique so two people can't share an email
      password_hash   — we NEVER store plain-text passwords; only the hash
      created_at      — ISO 8601 timestamp of when the account was made
      uploads_today   — how many transcriptions the user has done today
      last_upload_date— the date (YYYY-MM-DD) of the most recent upload,
                        used to know when to reset uploads_today back to 0
    """
    conn = get_db()
    # WAL mode is a persistent setting on the DB file — readers no longer block
    # the writer (and vice-versa), which matters under concurrent requests.
    conn.execute("PRAGMA journal_mode = WAL")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id               INTEGER PRIMARY KEY AUTOINCREMENT,
            email            TEXT UNIQUE NOT NULL,
            password_hash    TEXT NOT NULL,
            first_name       TEXT,
            last_name        TEXT,
            avatar           TEXT,
            created_at       TEXT,
            tier             TEXT DEFAULT 'free',  -- 'free', 'student', or 'pro'
            tokens_today     INTEGER DEFAULT 0,    -- words transcribed today
            last_token_date  TEXT,                 -- date of last transcription (YYYY-MM-DD)
            bonus_tokens     INTEGER DEFAULT 0,    -- extra daily tokens from referrals
            referral_code    TEXT UNIQUE,
            referred_by      INTEGER,
            is_admin         INTEGER DEFAULT 0     -- owner override: never limited
        )
    """)
    conn.commit()
    # Transcription history table — one row per transcription
    conn.execute("""
        CREATE TABLE IF NOT EXISTS transcriptions (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            text        TEXT NOT NULL,
            word_count  INTEGER DEFAULT 0,
            created_at  TEXT NOT NULL
        )
    """)
    conn.commit()

    # Password reset codes — one row per emailed code.
    # The code itself is hashed, never stored in the clear, so a leaked
    # database can't be used to take over accounts. Rows are consumed on use
    # and cleaned up as new ones are issued.
    conn.execute("""
        CREATE TABLE IF NOT EXISTS password_resets (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            code_hash   TEXT NOT NULL,
            expires_at  TEXT NOT NULL,   -- ISO timestamp (UTC)
            attempts    INTEGER DEFAULT 0,
            used        INTEGER DEFAULT 0,
            created_at  TEXT NOT NULL
        )
    """)
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_password_resets_user ON password_resets(user_id)"
    )
    conn.commit()

    # Login 2FA codes — same shape as password_resets, kept separate because
    # the two are semantically different (this one only ever grants a session,
    # never touches the password), so mixing them would risk one flow's fix
    # accidentally changing the other's behaviour.
    conn.execute("""
        CREATE TABLE IF NOT EXISTS login_otps (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            code_hash   TEXT NOT NULL,
            expires_at  TEXT NOT NULL,
            attempts    INTEGER DEFAULT 0,
            used        INTEGER DEFAULT 0,
            created_at  TEXT NOT NULL
        )
    """)
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_login_otps_user ON login_otps(user_id)"
    )
    conn.commit()

    # Feature requests — one row per submission from the profile dropdown.
    conn.execute("""
        CREATE TABLE IF NOT EXISTS feature_requests (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id     INTEGER NOT NULL,
            text        TEXT NOT NULL,
            created_at  TEXT NOT NULL
        )
    """)
    conn.commit()

    # Migrate any columns missing from older database versions
    migrations = [
        "first_name TEXT", "last_name TEXT", "avatar TEXT",
        "tier TEXT DEFAULT 'free'",
        "tokens_today INTEGER DEFAULT 0",
        "last_token_date TEXT",
        "bonus_tokens INTEGER DEFAULT 0",
        "referral_code TEXT", "referred_by INTEGER",
        "is_admin INTEGER DEFAULT 0",
    ]
    for col_def in migrations:
        try:
            conn.execute(f"ALTER TABLE users ADD COLUMN {col_def}")
            conn.commit()
        except Exception:
            pass

    # Add share_token to transcriptions if it doesn't exist yet
    try:
        conn.execute("ALTER TABLE transcriptions ADD COLUMN share_token TEXT")
        conn.commit()
    except Exception:
        pass

    # Fix any users whose tier column is NULL (inserted before migration added the column)
    conn.execute("UPDATE users SET tier = 'free' WHERE tier IS NULL")
    conn.commit()

    # Add title column so users can rename transcriptions
    try:
        conn.execute("ALTER TABLE transcriptions ADD COLUMN title TEXT")
        conn.commit()
    except Exception:
        pass

    # Add Google OAuth token columns
    for col in ["google_access_token TEXT", "google_refresh_token TEXT", "google_token_expiry TEXT"]:
        try:
            conn.execute(f"ALTER TABLE users ADD COLUMN {col}")
            conn.commit()
        except Exception:
            pass

    # Tracks whether a user has a real, self-chosen password. Accounts created
    # via "Continue with Google" get a random, never-shown password hash, so
    # password-confirmation flows (e.g. account deletion) must skip them.
    try:
        conn.execute("ALTER TABLE users ADD COLUMN has_password INTEGER DEFAULT 1")
        conn.commit()
    except Exception:
        pass

    # Stripe billing columns
    for col in ["stripe_customer_id TEXT", "stripe_subscription_id TEXT", "stripe_cancel_at_period_end INTEGER DEFAULT 0",
                "stripe_period_end INTEGER"]:
        try:
            conn.execute(f"ALTER TABLE users ADD COLUMN {col}")
            conn.commit()
        except Exception:
            pass

    # Notion OAuth columns
    for col in ["notion_access_token TEXT", "notion_workspace_name TEXT", "notion_bot_id TEXT"]:
        try:
            conn.execute(f"ALTER TABLE users ADD COLUMN {col}")
            conn.commit()
        except Exception:
            pass

    # Apple Sign-In — Apple's stable per-user identifier ("sub" claim)
    try:
        conn.execute("ALTER TABLE users ADD COLUMN apple_sub TEXT")
        conn.commit()
    except Exception:
        pass

    # Notebooks — user-created folders to organise transcriptions
    conn.execute("""
        CREATE TABLE IF NOT EXISTS notebooks (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id    INTEGER NOT NULL,
            name       TEXT NOT NULL,
            color      TEXT DEFAULT '#c9a96e',
            created_at TEXT NOT NULL
        )
    """)
    conn.commit()

    # Many-to-many: which transcriptions belong to which notebook
    conn.execute("""
        CREATE TABLE IF NOT EXISTS notebook_transcriptions (
            notebook_id      INTEGER NOT NULL,
            transcription_id INTEGER NOT NULL,
            PRIMARY KEY (notebook_id, transcription_id)
        )
    """)
    conn.commit()

    # Indexes on hot query paths
    conn.execute("CREATE INDEX IF NOT EXISTS idx_transcriptions_user ON transcriptions(user_id)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_transcriptions_user_id_desc ON transcriptions(user_id, id DESC)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_notebooks_user ON notebooks(user_id)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_users_email ON users(email)")
    conn.commit()
    conn.close()


# Run init_db() immediately when the module loads so the table always exists
# before any request can come in.
init_db()


# ── Auth decorator ─────────────────────────────────────────────────────────────

def login_required(f):
    """
    A decorator that protects a route so only logged-in users can access it.

    A decorator is a function that wraps another function to add behaviour.
    Here, before the real route function runs, we check whether "user_id" is
    stored in the session cookie. If it isn't, the user isn't logged in, so
    we send them to the login page instead.

    Usage:
        @app.route("/some-protected-page")
        @login_required          ← add this line right before the function
        def some_page():
            ...
    """
    @wraps(f)  # preserves the original function's name and docstring
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            # POST requests are always fetch/API calls in this app, so return
            # JSON — otherwise the browser gets an HTML redirect it can't parse.
            if request.method != "GET":
                return jsonify({"error": "Not logged in."}), 401
            return redirect(url_for("landing"))
        return f(*args, **kwargs)
    return decorated


def allowed_file(filename: str) -> bool:
    """Return True if the filename has an extension we accept."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# ── Routes ─────────────────────────────────────────────────────────────────────

# ── Auth routes ───────────────────────────────────────────────────────────────

@app.route("/login")
def landing():
    """
    GET /login  (also the landing page for non-logged-in users)
    Shows the landing + login/signup page (login.html has both panels).
    """
    if session.get("user_id"):
        return redirect(url_for("index"))
    return render_template("login.html")


@app.route("/privacy")
def privacy():
    """Public privacy policy — required for Google OAuth verification."""
    return render_template("privacy.html")


@app.route("/terms")
def terms():
    """Public terms of service — linked from the OAuth consent screen."""
    return render_template("terms.html")


@app.route("/favicon.ico")
def favicon():
    return app.send_static_file("favicon.png")

@app.route("/robots.txt")
def robots_txt():
    body = (
        "User-agent: *\n"
        "Allow: /\n"
        "Disallow: /history\n"
        "Disallow: /notebooks\n"
        "Disallow: /transcription/\n"
        "Disallow: /transcriptions/\n"
        f"Sitemap: {request.url_root.rstrip('/')}/sitemap.xml\n"
    )
    return Response(body, mimetype="text/plain")


@app.route("/sitemap.xml")
def sitemap_xml():
    base = request.url_root.rstrip("/")
    urls = ["/", "/login", "/privacy", "/terms"]
    items = "".join(
        f"<url><loc>{base}{u}</loc><changefreq>weekly</changefreq></url>" for u in urls
    )
    body = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
        f"{items}</urlset>"
    )
    return Response(body, mimetype="application/xml")

# Keep /login as an alias so old links still work


@app.route("/login", methods=["POST"])
@limiter.limit("10 per minute")
def login_post():
    """
    POST /login
    Expects a JSON body: { "email": "...", "password": "..." }

    We look the email up in the database. If found, we use Werkzeug's
    check_password_hash() to verify the password against the stored hash.
    On success we store the user's id and email in the session and tell
    the browser to redirect to the home page.

    Returns JSON so the front-end fetch() call can read the result:
      { "ok": true }          — success
      { "error": "..." }      — failure with a human-readable reason
    """
    data = request.get_json(silent=True)  # silent=True returns None on parse error

    # Make sure we actually received JSON with the fields we need
    if not data or not data.get("email") or not data.get("password"):
        return jsonify({"error": "Email and password are required."}), 400

    email    = data["email"].strip().lower()  # normalise so "User@Example.com" == "user@example.com"
    password = data["password"]

    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
    conn.close()

    # check_password_hash() returns False if the hash doesn't match, or if user is None
    if not user or not check_password_hash(user["password_hash"], password):
        # We give the same vague error for both "no account" and "wrong password"
        # so an attacker can't tell which emails are registered.
        return jsonify({"error": "Incorrect email or password."}), 401

    # Password is correct. On a browser we've already verified for this
    # account, that's enough — go straight in like before. Anywhere else,
    # the password alone doesn't create a session; a code has to follow.
    if is_trusted_device(user["id"]):
        session["user_id"]    = user["id"]
        session["user_email"] = user["email"]
        return jsonify({"ok": True})

    issue_login_otp(user)
    # Marks which login this code belongs to without granting access — every
    # route that actually does something still requires session["user_id"],
    # which is only set once /login/verify-otp succeeds.
    session["pending_2fa_user_id"] = user["id"]
    return jsonify({"otp_required": True, "email": user["email"]})


@app.route("/login/verify-otp", methods=["POST"])
@limiter.limit("10 per minute")
def login_verify_otp():
    """
    POST /login/verify-otp  { "code": "123456" }
    Completes a login that /login left pending on a new device. Trusts this
    browser for TRUSTED_DEVICE_DAYS afterward so the code isn't asked again
    on every visit — only when it's a browser we haven't verified before.
    """
    user_id = session.get("pending_2fa_user_id")
    if not user_id:
        return jsonify({"error": "Nothing to verify — please log in again."}), 400

    data = request.get_json(silent=True) or {}
    code = (data.get("code") or "").strip()
    if not code:
        return jsonify({"error": "Enter the code from your email."}), 400

    conn = get_db()
    user = conn.execute("SELECT id, email FROM users WHERE id = ?", (user_id,)).fetchone()
    if not user:
        conn.close()
        session.pop("pending_2fa_user_id", None)
        return jsonify({"error": "Nothing to verify — please log in again."}), 400

    row, err = _check_otp(conn, "login_otps", user_id, code, LOGIN_OTP_MAX_ATTEMPTS)
    conn.close()
    if err:
        return jsonify({"error": err}), 400

    session.pop("pending_2fa_user_id", None)
    session["user_id"]    = user["id"]
    session["user_email"] = user["email"]

    resp = jsonify({"ok": True})
    return mark_device_trusted(resp, user["id"])


@app.route("/login/resend-otp", methods=["POST"])
@limiter.limit("3 per minute; 10 per hour")
def login_resend_otp():
    """POST /login/resend-otp — issues a fresh code for the login left
    pending by /login, retiring whatever code came before it."""
    user_id = session.get("pending_2fa_user_id")
    if not user_id:
        return jsonify({"error": "Nothing to resend — please log in again."}), 400

    conn = get_db()
    user = conn.execute("SELECT id, email, first_name FROM users WHERE id = ?", (user_id,)).fetchone()
    conn.close()
    if not user:
        session.pop("pending_2fa_user_id", None)
        return jsonify({"error": "Nothing to resend — please log in again."}), 400

    issue_login_otp(user)
    return jsonify({"ok": True, "message": "A new code is on its way."})


@app.route("/signup", methods=["POST"])
@limiter.limit("5 per minute")
def signup_post():
    """
    POST /signup
    Expects a JSON body: { "email": "...", "password": "...", "confirm": "..." }

    We validate the inputs, hash the password with Werkzeug, insert a new row
    into the users table, then auto-login the user exactly like /login does.

    Returns:
      { "ok": true }     — account created and logged in
      { "error": "..." } — validation or DB error
    """
    data = request.get_json(silent=True)

    if not data or not data.get("email") or not data.get("password"):
        return jsonify({"error": "Email and password are required."}), 400

    email      = data["email"].strip().lower()
    password   = data["password"]
    confirm    = data.get("confirm", "")
    first_name   = data.get("first_name", "").strip()
    last_name    = data.get("last_name", "").strip()
    referral_in  = data.get("referral_code", "").strip().upper()  # code they were given

    if not first_name or not last_name:
        return jsonify({"error": "Please enter your first and last name."}), 400

    if len(password) < 6:
        return jsonify({"error": "Password must be at least 6 characters."}), 400

    if password != confirm:
        return jsonify({"error": "Passwords do not match."}), 400

    password_hash = generate_password_hash(password, method="pbkdf2:sha256")
    created_at    = datetime.utcnow().isoformat()

    # Generate a unique 8-character referral code for this new user.
    # secrets.token_urlsafe gives a random URL-safe string; we take 6 chars and uppercase it.
    new_ref_code = secrets.token_urlsafe(6).upper()[:8]

    # Look up who referred this new user (if anyone)
    conn = get_db()
    referrer_id = None
    if referral_in:
        referrer = conn.execute(
            "SELECT id FROM users WHERE referral_code = ?", (referral_in,)
        ).fetchone()
        if referrer:
            referrer_id = referrer["id"]

    try:
        cursor = conn.execute(
            """INSERT INTO users
               (email, password_hash, first_name, last_name, created_at, referral_code, referred_by, tier)
               VALUES (?, ?, ?, ?, ?, ?, ?, 'free')""",
            (email, password_hash, first_name, last_name, created_at, new_ref_code, referrer_id),
        )
        conn.commit()
        new_id = cursor.lastrowid
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({"error": "An account with that email already exists."}), 409

    # Reward the referrer with extra daily tokens for each person they bring in
    if referrer_id:
        conn.execute(
            "UPDATE users SET bonus_tokens = bonus_tokens + ? WHERE id = ?",
            (REFERRAL_BONUS_TOKENS, referrer_id),
        )
        conn.commit()

    conn.close()

    session["user_id"]    = new_id
    session["user_email"] = email

    # This browser just created the account, which is as verified as a device
    # gets — no reason to challenge its very next login with a 2FA code.
    resp = jsonify({"ok": True})
    return mark_device_trusted(resp, new_id)


# ── Password reset (emailed one-time code) ────────────────────────────────────
# Three steps: request a code, check it's right, then set the new password.
# The middle step exists only so the UI can advance before asking for a
# password — it grants nothing, and the final step re-verifies the code, so
# skipping straight to /reset-password is no easier.

def _check_otp(conn, table, user_id, code, max_attempts):
    """
    Return the matching live code row from `table` for `user_id`, or
    (None, error_message). Shared by password reset and login 2FA — both
    store a hashed, expiring, attempt-limited one-time code and differ only
    in what a successful check unlocks, which is entirely up to the caller.

    `table` is always one of our own two hardcoded literals, never request
    input, so building the query with an f-string here is safe.

    Counts a failed attempt against the newest outstanding code so guessing is
    bounded, and treats expired/used/burnt rows as simply not matching.
    """
    row = conn.execute(
        f"""SELECT * FROM {table}
            WHERE user_id = ? AND used = 0
            ORDER BY id DESC LIMIT 1""",
        (user_id,),
    ).fetchone()

    if not row:
        return None, "That code isn't valid. Request a new one."

    if datetime.utcnow() > datetime.fromisoformat(row["expires_at"]):
        return None, "That code has expired. Request a new one."

    if row["attempts"] >= max_attempts:
        return None, "Too many incorrect attempts. Request a new code."

    if not check_password_hash(row["code_hash"], code):
        conn.execute(f"UPDATE {table} SET attempts = attempts + 1 WHERE id = ?", (row["id"],))
        conn.commit()
        left = max_attempts - (row["attempts"] + 1)
        if left <= 0:
            return None, "Too many incorrect attempts. Request a new code."
        return None, f"That code isn't right. {left} attempt{'s' if left != 1 else ''} left."

    return row, None


def _find_user_by_email(conn, email):
    return conn.execute(
        "SELECT id, email, first_name FROM users WHERE email = ?", (email,)
    ).fetchone()


@app.route("/forgot-password", methods=["POST"])
@limiter.limit("3 per minute; 10 per hour")
def forgot_password():
    """
    POST /forgot-password  { "email": "..." }
    Emails a 6-digit code if the address has an account.

    Always responds the same way whether or not the account exists — a
    different response here would turn this into an endpoint for discovering
    which email addresses are registered.
    """
    data  = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()

    generic = jsonify({
        "ok": True,
        "message": "If that email has an account, a code is on its way.",
    })

    if not email:
        return jsonify({"error": "Enter your email address."}), 400

    conn = get_db()
    user = _find_user_by_email(conn, email)
    if not user:
        conn.close()
        return generic

    code = f"{secrets.randbelow(1_000_000):06d}"
    now  = datetime.utcnow()

    # Retire any outstanding codes so only the newest one works.
    conn.execute("UPDATE password_resets SET used = 1 WHERE user_id = ? AND used = 0", (user["id"],))
    conn.execute(
        """INSERT INTO password_resets (user_id, code_hash, expires_at, created_at)
           VALUES (?, ?, ?, ?)""",
        (
            user["id"],
            generate_password_hash(code, method="pbkdf2:sha256"),
            (now + timedelta(minutes=RESET_CODE_TTL_MINUTES)).isoformat(),
            now.isoformat(),
        ),
    )
    conn.commit()
    conn.close()

    name = user["first_name"] or "there"
    send_email(
        user["email"],
        "Your NoteCloud password reset code",
        f"""Hi {name},

Your NoteCloud password reset code is:

    {code}

It expires in {RESET_CODE_TTL_MINUTES} minutes and can only be used once.

If you didn't ask to reset your password you can ignore this email — your
current password still works and nothing has changed.

— NoteCloud
""",
    )
    return generic


@app.route("/verify-reset-code", methods=["POST"])
@limiter.limit("10 per minute")
def verify_reset_code():
    """
    POST /verify-reset-code  { "email": "...", "code": "123456" }
    Checks a code without consuming it, so the UI can show the new-password
    fields before committing. Grants nothing on its own.
    """
    data  = request.get_json(silent=True) or {}
    email = (data.get("email") or "").strip().lower()
    code  = (data.get("code") or "").strip()

    if not email or not code:
        return jsonify({"error": "Enter the code from your email."}), 400

    conn = get_db()
    user = _find_user_by_email(conn, email)
    if not user:
        conn.close()
        return jsonify({"error": "That code isn't valid. Request a new one."}), 400

    row, err = _check_otp(conn, "password_resets", user["id"], code, RESET_MAX_ATTEMPTS)
    conn.close()
    if err:
        return jsonify({"error": err}), 400
    return jsonify({"ok": True})


@app.route("/reset-password", methods=["POST"])
@limiter.limit("10 per minute")
def reset_password():
    """
    POST /reset-password
      { "email": "...", "code": "123456", "new_password": "...", "confirm": "..." }
    Verifies the code once more, sets the new password, and burns the code.
    """
    data     = request.get_json(silent=True) or {}
    email    = (data.get("email") or "").strip().lower()
    code     = (data.get("code") or "").strip()
    new_pw   = data.get("new_password") or ""
    confirm  = data.get("confirm") or ""

    if not email or not code:
        return jsonify({"error": "Enter the code from your email."}), 400
    if len(new_pw) < 6:
        return jsonify({"error": "Password must be at least 6 characters."}), 400
    if new_pw != confirm:
        return jsonify({"error": "Those passwords don't match."}), 400

    conn = get_db()
    user = _find_user_by_email(conn, email)
    if not user:
        conn.close()
        return jsonify({"error": "That code isn't valid. Request a new one."}), 400

    row, err = _check_otp(conn, "password_resets", user["id"], code, RESET_MAX_ATTEMPTS)
    if err:
        conn.close()
        return jsonify({"error": err}), 400

    # has_password is set because this is now a real, self-chosen password —
    # Google-created accounts start without one, and this is how they get one.
    conn.execute(
        "UPDATE users SET password_hash = ?, has_password = 1 WHERE id = ?",
        (generate_password_hash(new_pw, method="pbkdf2:sha256"), user["id"]),
    )
    conn.execute("UPDATE password_resets SET used = 1 WHERE id = ?", (row["id"],))
    conn.commit()
    conn.close()

    # Log out any existing sessions on this browser — whoever reset the
    # password should have to sign in with it.
    session.clear()

    return jsonify({"ok": True, "message": "Password updated. You can sign in now."})


@app.route("/logout")
def logout():
    """
    GET /logout
    Wipe the session (which removes user_id and user_email) and send the
    user back to the login page.
    """
    session.clear()
    return redirect(url_for("landing"))


# ── App routes ─────────────────────────────────────────────────────────────────

def get_token_status(user):
    """
    Given a user row, return a dict with their token usage info for today.
    This is used by both the index route and the transcribe route.

      limit       — total daily token budget (None = unlimited)
      used        — tokens used today
      remaining   — tokens left today (None = unlimited)
    """
    today = date.today().isoformat()
    tier  = user["tier"] or "free"
    used  = user["tokens_today"] if user["last_token_date"] == today else 0

    # Unlimited only for cases we've deliberately granted it: admins, the "dev"
    # tier, and tiers explicitly mapped to None (pro). An unrecognised tier
    # falls back to the free limit rather than to unlimited — failing open here
    # would hand out free unlimited access on any typo or stale tier value.
    if user["is_admin"] or tier == "dev":
        return {"limit": None, "used": used, "remaining": None, "tier": tier}

    if tier in TIER_LIMITS:
        base_limit = TIER_LIMITS[tier]
        if base_limit is None:      # pro
            return {"limit": None, "used": used, "remaining": None, "tier": tier}
    else:
        base_limit = TIER_LIMITS["free"]

    daily_limit = base_limit + (user["bonus_tokens"] or 0)
    return {
        "limit": daily_limit,
        "used": used,
        "remaining": max(0, daily_limit - used),
        "tier": tier,
    }


@app.route("/")
@login_required
def index():
    """Serve the main page, passing the logged-in user's info to the template."""
    conn = get_db()
    user = conn.execute(
        """SELECT first_name, last_name, email, avatar, referral_code,
                  bonus_tokens, tokens_today, last_token_date, tier, is_admin, has_password,
                  stripe_subscription_id, stripe_cancel_at_period_end, stripe_period_end
           FROM users WHERE id = ?""",
        (session["user_id"],)
    ).fetchone()
    conn.close()

    # If the user row is gone (e.g. DB was reset), clear the session and redirect
    if user is None:
        session.clear()
        return redirect(url_for("login"))

    status = get_token_status(user)

    # A cancelled-but-not-yet-expired plan still reads as paid in `tier`, so the
    # subscription panel needs these to tell "active" apart from "ending soon".
    has_stripe_sub = bool(user["stripe_subscription_id"])
    ending_soon = bool(user["stripe_cancel_at_period_end"])
    period_end  = user["stripe_period_end"]
    period_end_display = (
        datetime.fromtimestamp(period_end).strftime("%B %-d, %Y") if period_end else None
    )

    # Owner/dev grants (redeem_code sets is_admin=1 and tier='dev' together —
    # see the /redeem route) aren't billing subscriptions, so there is nothing
    # for a "Cancel subscription" button to do. Distinguish that from a paid
    # tier that was granted manually with no Stripe record behind it (e.g. a
    # comped account) — that one genuinely can self-downgrade to free, and
    # /cancel-subscription already supports it.
    grant_only = (not has_stripe_sub) and (user["is_admin"] or user["tier"] == "dev")

    return render_template(
        "index.html",
        user=user,
        tokens_remaining=status["remaining"],   # None = unlimited
        tokens_limit=status["limit"],
        tokens_used=status["used"],
        tier=status["tier"],
        referral_bonus_tokens=REFERRAL_BONUS_TOKENS,
        tier_limits=TIER_LIMITS,
        has_password=bool(user["has_password"]),
        ending_soon=ending_soon,
        period_end_display=period_end_display,
        has_stripe_sub=has_stripe_sub,
        grant_only=grant_only,
    )


@app.route("/redeem", methods=["POST"])
@login_required
@limiter.limit("5 per minute; 20 per hour")
def redeem_code():
    """
    POST /redeem  { "code": "..." }
    Checks the submitted code against the OWNER_CODE in .env.
    If it matches, grants is_admin = 1 (unlimited uploads forever).
    """
    data = request.get_json(silent=True)
    code = (data or {}).get("code", "").strip()

    if not code:
        return jsonify({"error": "Please enter a code."}), 400

    if OWNER_CODE and code.upper() == OWNER_CODE.upper():
        conn = get_db()
        conn.execute("UPDATE users SET is_admin = 1, tier = 'dev' WHERE id = ?", (session["user_id"],))
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "message": "Unlimited uploads unlocked!"})

    return jsonify({"error": "That code isn't valid."}), 400


@app.route("/feature-request", methods=["POST"])
@login_required
@limiter.limit("5 per hour")
def feature_request():
    """
    POST /feature-request  { "text": "..." }
    Stores a feature request and emails it to the site owner so it isn't
    only sitting in the database waiting to be queried.
    """
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()

    if not text:
        return jsonify({"error": "Enter what you'd like to see added."}), 400
    if len(text) > 2000:
        return jsonify({"error": "That's a bit long — please keep it under 2000 characters."}), 400

    conn = get_db()
    user = conn.execute(
        "SELECT email, first_name, last_name FROM users WHERE id = ?", (session["user_id"],)
    ).fetchone()
    conn.execute(
        "INSERT INTO feature_requests (user_id, text, created_at) VALUES (?, ?, ?)",
        (session["user_id"], text, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()

    who = f"{(user['first_name'] or '').strip()} {(user['last_name'] or '').strip()}".strip() or "A user"
    # Best-effort — the request is already saved in the DB either way, so a
    # failed notification email doesn't lose it, just delays you seeing it.
    send_email(
        EMAIL_USER or (user["email"] if user else ""),
        f"NoteCloud feature request from {who}",
        f"{who} ({user['email'] if user else 'unknown email'}) submitted a feature request:\n\n{text}",
    )

    return jsonify({"ok": True, "message": "Thanks — sent. We read every one of these."})


def _switch_plan(user_id, subscription_id, price_id, tier):
    """
    Move an existing subscription onto `price_id`, prorating the difference.

    Returns a Flask response to send back, or None if the stored subscription
    is no longer usable and the caller should fall back to a fresh checkout.

    The tier is written here rather than waiting for the webhook: unlike
    checkout — where the browser redirect proves nothing until Stripe confirms
    payment — this is our own authenticated API call, so a success response is
    Stripe confirming the change. The webhook still fires and is idempotent.
    """
    try:
        sub = stripe.Subscription.retrieve(subscription_id)
    except Exception as e:
        app.logger.warning("stripe: stored subscription %s unreadable: %s", subscription_id, e)
        return None

    if sub.get("status") not in ("active", "trialing", "past_due"):
        return None   # lapsed — let them start a new subscription

    try:
        item = sub["items"]["data"][0]
    except (KeyError, IndexError):
        return None

    if item["price"]["id"] == price_id and not sub.get("cancel_at_period_end"):
        return jsonify({"error": "You're already on that plan."}), 400

    try:
        updated = stripe.Subscription.modify(
            subscription_id,
            items=[{"id": item["id"], "price": price_id}],
            proration_behavior="create_prorations",
            cancel_at_period_end=False,   # switching plans also un-cancels
        )
    except Exception as e:
        app.logger.error("stripe plan switch error: %s", e)
        return jsonify({"error": "Could not change your plan — please try again."}), 500

    conn = get_db()
    conn.execute(
        """UPDATE users SET tier = ?, stripe_cancel_at_period_end = 0, stripe_period_end = ?
           WHERE id = ?""",
        (tier_from_subscription(updated) or tier, subscription_period_end(updated), user_id),
    )
    conn.commit()
    conn.close()

    return jsonify({
        "ok": True,
        "switched": True,
        "message": f"You're now on {tier.title()}. Any difference is prorated on your next invoice.",
    })


@app.route("/upgrade", methods=["POST"])
@login_required
def upgrade():
    """
    POST /upgrade  { "tier": "student" | "pro", "period": "monthly" | "annual" }
    Creates a Stripe Checkout Session and returns its URL. The user's tier is
    NOT upgraded here — that only happens once Stripe confirms payment via the
    /stripe/webhook route. Doing it here would let anyone hit this endpoint
    and grant themselves a paid tier for free.

    Needs in .env: STRIPE_SECRET_KEY and one Price ID per tier/period, e.g.
      STRIPE_PRICE_STUDENT_MONTHLY=price_...
      STRIPE_PRICE_STUDENT_ANNUAL=price_...
      STRIPE_PRICE_PRO_MONTHLY=price_...
      STRIPE_PRICE_PRO_ANNUAL=price_...
    """
    if not STRIPE_LIBS_AVAILABLE or not STRIPE_SECRET_KEY:
        return jsonify({"error": "coming_soon", "message": "Payments coming soon — stay tuned!"}), 501

    data = request.get_json(silent=True) or {}
    tier   = data.get("tier", "pro")
    period = data.get("period", "monthly")
    # The pricing UI's billing toggle uses "yearly" — treat it as "annual".
    if period == "yearly":
        period = "annual"
    if tier not in ("student", "pro"):
        return jsonify({"error": "Invalid tier."}), 400
    if period not in ("monthly", "annual"):
        return jsonify({"error": "Invalid billing period."}), 400

    price_id = STRIPE_PRICE_IDS.get((tier, period))
    if not price_id:
        return jsonify({"error": "That plan isn't configured yet."}), 503

    user_id = session["user_id"]
    conn = get_db()
    user = conn.execute(
        "SELECT email, stripe_customer_id, stripe_subscription_id FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()
    conn.close()

    # Already subscribed? Move the existing subscription onto the new price
    # rather than opening a second checkout. Starting a fresh subscription here
    # would leave the old one live and billing alongside it, and the webhook
    # would overwrite stripe_subscription_id — orphaning a charge the user can
    # no longer see or cancel from the UI.
    if user["stripe_subscription_id"]:
        switched = _switch_plan(user_id, user["stripe_subscription_id"], price_id, tier)
        if switched is not None:
            return switched
        # Falls through to checkout when the stored subscription is already
        # gone on Stripe's side (expired, or refunded and deleted).

    try:
        checkout_kwargs = {
            "payment_method_types": ["card"],
            "line_items": [{"price": price_id, "quantity": 1}],
            "mode": "subscription",
            "client_reference_id": str(session["user_id"]),
            "success_url": request.host_url + "?checkout=success",
            "cancel_url": request.host_url,
        }
        if user["stripe_customer_id"]:
            checkout_kwargs["customer"] = user["stripe_customer_id"]
        else:
            checkout_kwargs["customer_email"] = user["email"]
        checkout = stripe.checkout.Session.create(**checkout_kwargs)
        return jsonify({"url": checkout.url})
    except Exception as e:
        app.logger.error("stripe checkout error: %s", e)
        return jsonify({"error": "Could not start checkout — please try again."}), 500


@app.route("/stripe/webhook", methods=["POST"])
@csrf.exempt  # Stripe posts here directly — no CSRF token to send. The Stripe
              # signature check inside this view is what verifies authenticity.
def stripe_webhook():
    """
    POST /stripe/webhook
    Receives billing lifecycle events from Stripe. This is the ONLY place a
    user's tier is granted or revoked based on payment — never from a redirect
    the browser controls, since that could be replayed or hit directly.

    Configure this URL in the Stripe Dashboard → Developers → Webhooks, and
    put the signing secret it gives you in STRIPE_WEBHOOK_SECRET. Every event
    MUST be signature-verified — an unsigned fallback here would let anyone
    POST a forged checkout.session.completed with an arbitrary
    client_reference_id and grant themselves Pro for free.
    """
    if not STRIPE_LIBS_AVAILABLE or not STRIPE_SECRET_KEY:
        return jsonify({"error": "Stripe not configured."}), 503

    if not STRIPE_WEBHOOK_SECRET:
        # A live Price ID doesn't imply a webhook is registered yet (that
        # secret is only issued once the endpoint is created in the Stripe
        # Dashboard, or via `stripe listen` for local testing) — refuse
        # rather than fall back to parsing an unverified body.
        app.logger.error("stripe webhook hit with STRIPE_WEBHOOK_SECRET unset — refusing unsigned event")
        return jsonify({"error": "Webhook not configured."}), 503

    payload    = request.get_data()
    sig_header = request.headers.get("Stripe-Signature", "")

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except Exception as e:
        app.logger.warning("stripe webhook signature/parse error: %s", e)
        return jsonify({"error": "Invalid payload."}), 400

    etype = event["type"]
    obj   = event["data"]["object"]

    conn = get_db()
    try:
        if etype == "checkout.session.completed":
            user_id = obj.get("client_reference_id")
            customer_id     = obj.get("customer")
            subscription_id = obj.get("subscription")
            if user_id:
                tier = "pro"
                try:
                    if subscription_id:
                        sub = stripe.Subscription.retrieve(subscription_id)
                        price_id = sub["items"]["data"][0]["price"]["id"]
                        tier = STRIPE_PRICE_TO_TIER.get(price_id, "pro")
                except Exception as e:
                    app.logger.warning("stripe: could not resolve tier from subscription: %s", e)
                conn.execute(
                    """UPDATE users SET tier = ?, stripe_customer_id = ?, stripe_subscription_id = ?,
                       stripe_cancel_at_period_end = 0 WHERE id = ?""",
                    (tier, customer_id, subscription_id, user_id),
                )
                conn.commit()

        elif etype == "customer.subscription.updated":
            sub_id = obj.get("id")
            cancel_at_period_end = 1 if obj.get("cancel_at_period_end") else 0
            # Re-derive the tier from whatever price the subscription now
            # carries — this event is what reports a plan switch, so without it
            # a Student→Pro change would bill correctly but never actually
            # grant Pro. Left alone if the price isn't one we recognise.
            new_tier = tier_from_subscription(obj)
            if new_tier:
                conn.execute(
                    """UPDATE users SET tier = ?, stripe_cancel_at_period_end = ?,
                       stripe_period_end = ? WHERE stripe_subscription_id = ?""",
                    (new_tier, cancel_at_period_end, subscription_period_end(obj), sub_id),
                )
            else:
                conn.execute(
                    """UPDATE users SET stripe_cancel_at_period_end = ?, stripe_period_end = ?
                       WHERE stripe_subscription_id = ?""",
                    (cancel_at_period_end, subscription_period_end(obj), sub_id),
                )
            conn.commit()

        elif etype == "customer.subscription.deleted":
            sub_id = obj.get("id")
            conn.execute(
                """UPDATE users SET tier = 'free', stripe_subscription_id = NULL,
                   stripe_cancel_at_period_end = 0, stripe_period_end = NULL
                   WHERE stripe_subscription_id = ?""",
                (sub_id,),
            )
            conn.commit()
    finally:
        conn.close()

    return jsonify({"ok": True})


@app.route("/history")
@login_required
def history():
    """
    GET /history
    Returns the last 50 transcriptions for the logged-in user as JSON,
    newest first. Also returns which notebook(s) each item belongs to.
    The sidebar calls this with fetch() to populate itself.
    """
    conn = get_db()
    rows = conn.execute(
        """SELECT id, text, word_count, created_at, title
           FROM transcriptions WHERE user_id = ?
           ORDER BY id DESC LIMIT 50""",
        (session["user_id"],)
    ).fetchall()
    items = [dict(r) for r in rows]

    # Attach notebook membership so the frontend can show a badge
    if items:
        trans_ids = [i["id"] for i in items]
        placeholders = ",".join("?" * len(trans_ids))
        nb_rows = conn.execute(
            f"SELECT notebook_id, transcription_id FROM notebook_transcriptions"
            f" WHERE transcription_id IN ({placeholders})",
            trans_ids
        ).fetchall()
        nb_map = {}
        for nr in nb_rows:
            nb_map.setdefault(nr["transcription_id"], []).append(nr["notebook_id"])
        for item in items:
            item["notebook_ids"] = nb_map.get(item["id"], [])

    conn.close()
    return jsonify({"items": items})


@app.route("/transcriptions/<int:trans_id>", methods=["DELETE"])
@login_required
def delete_transcription(trans_id):
    """DELETE /transcriptions/<id> — delete a single transcription."""
    conn = get_db()
    row = conn.execute(
        "SELECT id FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Not found."}), 404
    conn.execute("DELETE FROM notebook_transcriptions WHERE transcription_id = ?", (trans_id,))
    conn.execute("DELETE FROM transcriptions WHERE id = ?", (trans_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/history/clear", methods=["POST"])
@login_required
def clear_history():
    """DELETE all transcriptions for the logged-in user."""
    conn = get_db()
    conn.execute("DELETE FROM notebook_transcriptions WHERE transcription_id IN "
                 "(SELECT id FROM transcriptions WHERE user_id = ?)", (session["user_id"],))
    conn.execute("DELETE FROM transcriptions WHERE user_id = ?", (session["user_id"],))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/transcription/<int:trans_id>")
@login_required
def transcription_detail(trans_id):
    """Render the detail page for a single transcription."""
    conn = get_db()
    row = conn.execute(
        "SELECT id, text, word_count, created_at, title FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    user = conn.execute("SELECT tier, is_admin FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    conn.close()
    if not row:
        return "Not found", 404
    is_pro = bool(user and (user["is_admin"] or user["tier"] in ("pro", "dev")))
    return render_template("transcription.html", t=dict(row), is_pro=is_pro)


@app.route("/transcriptions/<int:trans_id>/rewrite", methods=["POST"])
@login_required
def rewrite_transcription(trans_id):
    """
    POST /transcriptions/<id>/rewrite  { "text": "...", "style": "longer|shorter|casual|professional|..." }
    Rewrites the transcription text using Claude and returns the new text.
    """
    err = require_pro_tier("AI Rewrite is available on the Pro plan.")
    if err: return err

    conn = get_db()
    row = conn.execute(
        "SELECT id FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    conn.close()
    if not row:
        return jsonify({"error": "Not found"}), 404

    data  = request.get_json(silent=True) or {}
    text  = (data.get("text") or "").strip()
    style = (data.get("style") or "").strip().lower()
    if not text or not style:
        return jsonify({"error": "Missing text or style"}), 400

    prompts = {
        "longer":       "Expand this text significantly, adding more detail and depth while keeping the same meaning:",
        "shorter":      "Condense this text to its essential points, keeping it clear and readable:",
        "casual":       "Rewrite this in a casual, conversational tone — like texting a friend:",
        "professional": "Rewrite this in a polished, professional tone suitable for a work email or report:",
        "bullets":      "Convert this into a clean bullet-point list, preserving all key information:",
        "simplify":     "Rewrite this in simple, plain language that anyone can understand:",
        "formal":       "Rewrite this in formal academic or business language:",
        "grammar":      "Fix all grammar, punctuation, and spelling errors in this text without changing the meaning:",
    }
    instruction = prompts.get(style, f"Rewrite this text to be more {style}:")

    client = _anthropic
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": f"{instruction}\n\n{text}"
        }]
    )
    rewritten = message.content[0].text if message.content else text
    return jsonify({"ok": True, "text": rewritten})


@app.route("/transcriptions/<int:trans_id>/save", methods=["POST"])
@login_required
def save_transcription(trans_id):
    """POST /transcriptions/<id>/save  { "text": "..." }  — update stored text."""
    conn = get_db()
    row = conn.execute(
        "SELECT id FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Not found"}), 404
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    word_count = len(text.split()) if text else 0
    conn.execute(
        "UPDATE transcriptions SET text = ?, word_count = ? WHERE id = ?",
        (text, word_count, trans_id)
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "word_count": word_count})


def require_pro_tier(message="This feature is available on the Pro plan."):
    """Return a JSON error response unless the user is on the pro (or admin/dev) tier."""
    conn = get_db()
    user = conn.execute("SELECT tier, is_admin FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    conn.close()
    if user and (user["is_admin"] or user["tier"] in ("pro", "dev")):
        return None
    return jsonify({"error": "upgrade_required", "message": message}), 403


def require_paid_tier(message="This feature is available on Student and Pro plans."):
    """Return a JSON error response if the user is on the free tier, else None.

    This is the ONLY thing that actually enforces a paid feature — hiding a
    button in the template is cosmetic and can be bypassed (disable JS, edit the
    DOM, or just POST to the endpoint directly). Every paid route must call this.
    """
    conn = get_db()
    user = conn.execute("SELECT tier, is_admin FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    conn.close()
    if user and (user["is_admin"] or user["tier"] not in ("free", None)):
        return None
    return jsonify({"error": "upgrade_required", "message": message}), 403


@app.route("/notebooks", methods=["GET"])
@login_required
def list_notebooks():
    err = require_paid_tier("Notebooks are available on Student and Pro plans.")
    if err: return err
    conn = get_db()
    rows = conn.execute(
        """SELECT n.id, n.name, n.color, n.created_at,
                  COUNT(nt.transcription_id) AS item_count
           FROM notebooks n
           LEFT JOIN notebook_transcriptions nt ON nt.notebook_id = n.id
           WHERE n.user_id = ?
           GROUP BY n.id
           ORDER BY n.id DESC""",
        (session["user_id"],)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/notebooks", methods=["POST"])
@login_required
def create_notebook():
    err = require_paid_tier("Notebooks are available on Student and Pro plans.")
    if err: return err
    data = request.get_json(silent=True) or {}
    name  = data.get("name", "").strip()
    color = data.get("color", "#c9a96e").strip()
    if not re.match(r'^#[0-9a-fA-F]{6}$', color):
        color = "#c9a96e"
    if not name:
        return jsonify({"error": "Notebook name is required."}), 400
    conn = get_db()
    cursor = conn.execute(
        "INSERT INTO notebooks (user_id, name, color, created_at) VALUES (?, ?, ?, ?)",
        (session["user_id"], name, color, datetime.utcnow().isoformat())
    )
    conn.commit()
    nb_id = cursor.lastrowid
    conn.close()
    return jsonify({"ok": True, "id": nb_id, "name": name, "color": color, "item_count": 0})


@app.route("/notebooks/<int:nb_id>", methods=["DELETE"])
@login_required
def delete_notebook(nb_id):
    """
    DELETE /notebooks/<id>
    Deletes the notebook and removes all its transcription assignments.
    The transcriptions themselves are NOT deleted.
    """
    conn = get_db()
    nb = conn.execute(
        "SELECT id FROM notebooks WHERE id = ? AND user_id = ?",
        (nb_id, session["user_id"])
    ).fetchone()
    if not nb:
        conn.close()
        return jsonify({"error": "Not found."}), 404
    conn.execute("DELETE FROM notebook_transcriptions WHERE notebook_id = ?", (nb_id,))
    conn.execute("DELETE FROM notebooks WHERE id = ?", (nb_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/notebooks/<int:nb_id>/items", methods=["GET"])
@login_required
def notebook_items(nb_id):
    """
    GET /notebooks/<id>/items
    Returns the transcriptions inside a notebook, newest first.
    """
    conn = get_db()
    nb = conn.execute(
        "SELECT id FROM notebooks WHERE id = ? AND user_id = ?",
        (nb_id, session["user_id"])
    ).fetchone()
    if not nb:
        conn.close()
        return jsonify({"error": "Not found."}), 404
    rows = conn.execute(
        """SELECT t.id, t.text, t.word_count, t.created_at
           FROM transcriptions t
           JOIN notebook_transcriptions nt ON nt.transcription_id = t.id
           WHERE nt.notebook_id = ?
           ORDER BY t.id DESC""",
        (nb_id,)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/notebooks/<int:nb_id>/items", methods=["POST"])
@login_required
def add_to_notebook(nb_id):
    """
    POST /notebooks/<id>/items  { "transcription_id": 123 }
    Adds a transcription to a notebook. Safe to call twice (idempotent).
    """
    data = request.get_json(silent=True) or {}
    trans_id = data.get("transcription_id")
    if not trans_id:
        return jsonify({"error": "transcription_id required."}), 400
    conn = get_db()
    nb = conn.execute(
        "SELECT id FROM notebooks WHERE id = ? AND user_id = ?",
        (nb_id, session["user_id"])
    ).fetchone()
    if not nb:
        conn.close()
        return jsonify({"error": "Notebook not found."}), 404
    trans = conn.execute(
        "SELECT id FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    if not trans:
        conn.close()
        return jsonify({"error": "Transcription not found."}), 404
    try:
        conn.execute(
            "INSERT INTO notebook_transcriptions (notebook_id, transcription_id) VALUES (?, ?)",
            (nb_id, trans_id)
        )
        conn.commit()
    except sqlite3.IntegrityError:
        pass  # already in notebook — that's fine
    conn.close()
    return jsonify({"ok": True})


@app.route("/notebooks/<int:nb_id>/items/<int:trans_id>", methods=["DELETE"])
@login_required
def remove_from_notebook(nb_id, trans_id):
    """
    DELETE /notebooks/<nb_id>/items/<trans_id>
    Removes a transcription from a notebook (doesn't delete the transcription itself).
    """
    conn = get_db()
    nb = conn.execute(
        "SELECT id FROM notebooks WHERE id = ? AND user_id = ?",
        (nb_id, session["user_id"])
    ).fetchone()
    if not nb:
        conn.close()
        return jsonify({"error": "Not found."}), 404
    conn.execute(
        "DELETE FROM notebook_transcriptions WHERE notebook_id = ? AND transcription_id = ?",
        (nb_id, trans_id)
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/transcriptions/<int:trans_id>/share", methods=["POST"])
@login_required
def share_transcription(trans_id):
    """
    POST /transcriptions/<id>/share
    Generates (or returns the existing) share token for a transcription.
    Only the owner of the transcription can call this.
    Returns { "url": "https://..." } with the public share link.
    """
    conn = get_db()
    row = conn.execute(
        "SELECT id, share_token FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()

    if not row:
        conn.close()
        return jsonify({"error": "Not found."}), 404

    token = row["share_token"]
    if not token:
        token = secrets.token_urlsafe(12)  # 16-char URL-safe string
        conn.execute(
            "UPDATE transcriptions SET share_token = ? WHERE id = ?",
            (token, trans_id)
        )
        conn.commit()
    conn.close()

    share_url = request.host_url.rstrip("/") + f"/s/{token}"
    return jsonify({"ok": True, "url": share_url})


@app.route("/s/<token>")
def view_shared(token):
    """
    GET /s/<token>
    Public page — no login required. Displays the shared transcription.
    """
    conn = get_db()
    row = conn.execute(
        """SELECT t.text, t.word_count, t.created_at,
                  u.first_name, u.last_name
           FROM transcriptions t
           JOIN users u ON u.id = t.user_id
           WHERE t.share_token = ?""",
        (token,)
    ).fetchone()
    conn.close()

    if not row:
        return "This link is invalid or has been removed.", 404

    return render_template("share.html", item=dict(row))


@app.route("/cancel-subscription", methods=["POST"])
@login_required
def cancel_subscription():
    """
    POST /cancel-subscription
    Cancels the user's paid plan. If they have a real Stripe subscription, it's
    set to cancel at the end of the current billing period (so they keep
    access they already paid for) — the actual tier downgrade to 'free' only
    happens later via the customer.subscription.deleted webhook, never here,
    so cancelling can't be used to keep paid access indefinitely.

    Accounts without a Stripe subscription (e.g. an admin/dev-code grant) are
    downgraded immediately since there's no billing period to honor.
    """
    user_id = session["user_id"]
    conn = get_db()
    user = conn.execute(
        "SELECT tier, is_admin, stripe_subscription_id FROM users WHERE id = ?", (user_id,)
    ).fetchone()

    # A real Stripe subscription is always cancellable, checked before the
    # grant-based rejection below. Being made an admin doesn't stop a card
    # being charged, so an admin who genuinely subscribed must still be able
    # to stop the billing — otherwise their only route out is the Stripe
    # dashboard, and the app would keep showing them as subscribed.
    if user["stripe_subscription_id"] and STRIPE_LIBS_AVAILABLE and STRIPE_SECRET_KEY:
        try:
            sub = stripe.Subscription.modify(
                user["stripe_subscription_id"], cancel_at_period_end=True
            )
            period_end = subscription_period_end(sub)
            conn.execute(
                "UPDATE users SET stripe_cancel_at_period_end = 1, stripe_period_end = ? WHERE id = ?",
                (period_end, user_id),
            )
            conn.commit()
            conn.close()
            return jsonify({
                "ok": True,
                "cancel_at_period_end": True,
                "period_end": period_end,
                "message": "Your plan will end at the current billing period — you'll keep access until then.",
            })
        except Exception as e:
            app.logger.error("stripe cancel error: %s", e)
            conn.close()
            return jsonify({"error": "Could not cancel — please try again."}), 500

    # Nothing billable on file. Admin and dev grants aren't subscriptions, and
    # free has nothing to cancel, so there's genuinely nothing to do here.
    if user["is_admin"] or user["tier"] in ("free", "dev", None):
        conn.close()
        return jsonify({"error": "No active subscription to cancel."}), 400

    # A paid tier granted manually (no Stripe record) — nothing to bill, so
    # downgrade right away.
    conn.execute("UPDATE users SET tier = 'free' WHERE id = ?", (user_id,))
    conn.commit()
    conn.close()

    return jsonify({"ok": True, "message": "Subscription cancelled. You've been moved to the free plan."})


@app.route("/resume-subscription", methods=["POST"])
@login_required
def resume_subscription():
    """
    POST /resume-subscription
    Undoes a pending cancellation — the subscription is still live until the
    period ends, so clearing cancel_at_period_end simply lets it renew as
    normal. Only valid while the plan is in the "ending soon" state; once the
    period actually lapses Stripe deletes the subscription and the user has to
    go through checkout again.
    """
    user_id = session["user_id"]
    conn = get_db()
    user = conn.execute(
        "SELECT stripe_subscription_id, stripe_cancel_at_period_end FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()

    if not user["stripe_subscription_id"] or not user["stripe_cancel_at_period_end"]:
        conn.close()
        return jsonify({"error": "No pending cancellation to undo."}), 400

    if not (STRIPE_LIBS_AVAILABLE and STRIPE_SECRET_KEY):
        conn.close()
        return jsonify({"error": "Stripe not configured."}), 503

    try:
        sub = stripe.Subscription.modify(
            user["stripe_subscription_id"], cancel_at_period_end=False
        )
        conn.execute(
            "UPDATE users SET stripe_cancel_at_period_end = 0, stripe_period_end = ? WHERE id = ?",
            (subscription_period_end(sub), user_id),
        )
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "message": "Your plan will keep renewing — welcome back."})
    except Exception as e:
        app.logger.error("stripe resume error: %s", e)
        conn.close()
        return jsonify({"error": "Could not resume — please try again."}), 500


@app.route("/profile/upload", methods=["POST"])
@login_required
def upload_avatar():
    """
    POST /profile/upload
    Accepts a profile picture upload, saves it to static/avatars/,
    and stores the filename in the database.
    """
    if "avatar" not in request.files:
        return jsonify({"error": "No file uploaded."}), 400

    file = request.files["avatar"]
    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file."}), 400

    if PIL_AVAILABLE:
        try:
            raw = file.read()
            PilImage.open(io.BytesIO(raw)).verify()
            file.stream.seek(0)
        except Exception:
            return jsonify({"error": "Invalid or corrupted image."}), 415

    # Save into static/avatars/ using the user's id as the filename
    # so each upload overwrites the previous one cleanly.
    ext = file.filename.rsplit(".", 1)[1].lower()
    filename = f"avatar_{session['user_id']}.{ext}"
    avatars_dir = os.path.join(os.path.dirname(__file__), "static", "avatars")
    os.makedirs(avatars_dir, exist_ok=True)
    file.save(os.path.join(avatars_dir, filename))

    conn = get_db()
    conn.execute("UPDATE users SET avatar = ? WHERE id = ?", (filename, session["user_id"]))
    conn.commit()
    conn.close()

    return jsonify({"ok": True, "avatar": filename})


@app.route("/transcribe", methods=["POST"])
@login_required  # ← only logged-in users can transcribe
def transcribe():
    """
    POST /transcribe
    Expects a multipart form upload with a field named "image".
    Returns JSON: { "transcription": "..." } on success
                  { "error": "..." }          on failure

    After a successful transcription we update the user's upload count for
    today. If their last upload was on a previous day we reset the counter
    to 1 (this fresh day's first upload). This tracking is the groundwork
    for enforcing FREE_DAILY_LIMIT in a future version.
    """

    # ── 0. Check the user's daily token limit ──────────────────────────────
    user_id   = session["user_id"]
    today_str = date.today().isoformat()

    conn = get_db()
    user = conn.execute(
        "SELECT tokens_today, last_token_date, bonus_tokens, tier, is_admin FROM users WHERE id = ?",
        (user_id,)
    ).fetchone()
    conn.close()

    status = get_token_status(user)
    if status["remaining"] is not None and status["remaining"] <= 0:
        return jsonify({
            "error": "limit_reached",
            "message": f"You've used all {status['limit']} tokens for today. Upgrade for more, or share your referral code to earn bonus tokens."
        }), 429

    # ── 1. Validate the upload(s) ───────────────────────────────────────────
    # Accept either a list of files (multi-page: images[]) or a single file (image)

    files = request.files.getlist("images[]")
    if not files or files[0].filename == "":
        # Fall back to legacy single-image field
        single = request.files.get("image")
        if not single or single.filename == "":
            return jsonify({"error": "No image field in the request."}), 400
        files = [single]

    # ── Enforce pages-per-upload by tier (server-side) ──────────────────────
    # The frontend also limits this, but that's cosmetic — a user could POST
    # more images directly, so the real gate lives here.
    tier = (user["tier"] if user else None) or "free"
    page_cap = None if (user and user["is_admin"]) else TIER_PAGE_LIMITS.get(tier)
    if page_cap is not None and len(files) > page_cap:
        return jsonify({
            "error": "page_limit",
            "message": (
                f"Your plan allows {page_cap} page{'s' if page_cap != 1 else ''} per upload. "
                "Upgrade to Student or Pro to transcribe more pages at once."
            )
        }), 403

    for f in files:
        if not allowed_file(f.filename):
            return jsonify({"error": f"Unsupported file type: {f.filename}. Use PNG, JPG, WEBP, or GIF."}), 415
        if PIL_AVAILABLE:
            try:
                raw = f.read()
                PilImage.open(io.BytesIO(raw)).verify()
                f.stream.seek(0)
            except Exception:
                return jsonify({"error": f"Invalid or corrupted image: {f.filename}"}), 415

    # ── 2. Read every image and convert to base64 ───────────────────────────

    image_blocks = []
    for f in files:
        extension = f.filename.rsplit(".", 1)[1].lower()
        mime_type = MIME_TYPES[extension]
        image_data = base64.standard_b64encode(f.read()).decode("utf-8")
        image_blocks.append({
            "type": "image",
            "source": {"type": "base64", "media_type": mime_type, "data": image_data},
        })

    # ── 3. Call the Anthropic API ───────────────────────────────────────────

    force = request.form.get("force") == "true"
    page_note = f" There are {len(files)} pages — transcribe them in order, separating pages with '---'." if len(files) > 1 else ""
    if force:
        prompt_text = (
            "Please transcribe all text visible in this image, whether handwritten, typed, or printed."
            + page_note
            + " Output only the transcribed text — no explanations, no formatting labels, "
            "no extra commentary. If you cannot read a word clearly, indicate it with [illegible]."
        )
    else:
        prompt_text = (
            "Please transcribe all of the handwritten text in this image."
            + page_note
            + " If the image contains no handwritten text (e.g. it shows typed, printed, or digital text only),"
            " respond with exactly: [NOT_HANDWRITTEN]"
            " Otherwise output only the transcribed text — no explanations, no formatting labels, "
            "no extra commentary. If you cannot read a word clearly, indicate it with [illegible]."
        )

    client = _anthropic

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": image_blocks + [{"type": "text", "text": prompt_text}],
                }
            ],
        )
    except anthropic.AuthenticationError:
        return jsonify({"error": "Invalid API key — check your ANTHROPIC_API_KEY in .env"}), 500
    except anthropic.APIError as e:
        return jsonify({"error": f"Claude API error: {e}"}), 500

    # The response is a list of content blocks; we want the first text block.
    transcription = message.content[0].text.strip()

    # Detect non-handwritten images and surface a confirmation to the user.
    if not force and transcription == "[NOT_HANDWRITTEN]":
        return jsonify({"error": "no_handwriting"}), 200

    # ── 4. Count words and update the user's token balance ─────────────────
    # Each word in the transcription costs one token.
    word_count = len(transcription.split())

    # Use BEGIN IMMEDIATE to hold a write lock while we re-check the limit
    # and apply the update atomically, preventing race conditions.
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("BEGIN IMMEDIATE")
    user = conn.execute(
        "SELECT tokens_today, last_token_date, bonus_tokens, tier, is_admin FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()

    fresh_status = get_token_status(user)

    # The limit was already checked before the transcription ran, so by this
    # point the work is done and the API call is spent. If the page turned out
    # to be longer than the remaining budget we still hand it over and charge
    # up to the cap — discarding it would mean the user loses the result AND
    # we've burned the call, and a page denser than the daily limit could never
    # be transcribed at all however many times they retried. The cap still
    # holds because the pre-check blocks the *next* upload at zero remaining.
    # Charged separately from word_count so history keeps the true length.
    tokens_charged = word_count
    if fresh_status["remaining"] is not None:
        tokens_charged = min(word_count, fresh_status["remaining"])

    if user["last_token_date"] == today_str:
        new_total = user["tokens_today"] + tokens_charged
    else:
        new_total = tokens_charged  # new day — reset

    conn.execute(
        "UPDATE users SET tokens_today = ?, last_token_date = ? WHERE id = ?",
        (new_total, today_str, user_id),
    )
    # Save to history
    cursor = conn.execute(
        "INSERT INTO transcriptions (user_id, text, word_count, created_at) VALUES (?, ?, ?, ?)",
        (user_id, transcription, word_count, datetime.utcnow().isoformat()),
    )
    transcription_id = cursor.lastrowid
    conn.commit()

    # Recalculate remaining for the frontend counter
    updated_user = conn.execute(
        "SELECT tokens_today, last_token_date, bonus_tokens, tier, is_admin FROM users WHERE id = ?",
        (user_id,)
    ).fetchone()
    conn.close()

    new_status = get_token_status(updated_user)

    return jsonify({
        "transcription": transcription,
        "transcription_id": transcription_id,
        "tokens_used": tokens_charged,
        "tokens_remaining": new_status["remaining"],  # None = unlimited
        "tokens_limit": new_status["limit"],
    })


@app.route("/cleanup", methods=["POST"])
@login_required
def cleanup_text():
    """
    POST /cleanup  { "text": "..." }
    Sends the transcribed text back to Claude to fix grammar, spelling,
    punctuation and formatting while keeping the meaning identical.
    """
    err = require_paid_tier("AI cleanup is available on Student and Pro plans.")
    if err: return err

    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "No text provided."}), 400

    client = _anthropic
    try:
        message = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4096,
            messages=[{
                "role": "user",
                "content": (
                    "Clean up the following transcribed handwriting. Fix grammar, spelling, "
                    "punctuation and spacing. Keep the exact same meaning and content — "
                    "just make it polished and readable. Output only the cleaned text.\n\n"
                    + text
                ),
            }],
        )
    except anthropic.APIError as e:
        app.logger.error("cleanup APIError: %s", e)
        return jsonify({"error": "Cleanup failed — please try again."}), 500

    return jsonify({"text": message.content[0].text})


@app.route("/transcriptions/<int:trans_id>/rename", methods=["POST"])
@login_required
def rename_transcription(trans_id):
    """
    POST /transcriptions/<id>/rename  { "title": "My custom name" }
    Sets a display title on a transcription (shown in the history sidebar).
    """
    data = request.get_json(silent=True) or {}
    title = data.get("title", "").strip()
    if not title:
        return jsonify({"error": "Title required."}), 400
    conn = get_db()
    row = conn.execute(
        "SELECT id FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Not found."}), 404
    conn.execute("UPDATE transcriptions SET title = ? WHERE id = ?", (title, trans_id))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ── Download routes ───────────────────────────────────────────────────────────

@app.route("/transcriptions/<int:trans_id>/download/txt")
@login_required
def download_txt(trans_id):
    """
    GET /transcriptions/<id>/download/txt
    Sends the transcription as a plain-text file download.
    send_file() tells the browser to download rather than display the content.
    """
    conn = get_db()
    row = conn.execute(
        "SELECT text, created_at FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    conn.close()

    if not row:
        return jsonify({"error": "Not found."}), 404

    # Wrap the text in a BytesIO buffer so we don't need to write a temp file
    buf = io.BytesIO(row["text"].encode("utf-8"))
    buf.seek(0)

    # Build a filename from the date, e.g. "note-cloud-2024-03-15.txt"
    date_part = (row["created_at"] or "")[:10]  # grab YYYY-MM-DD
    filename = f"note-cloud-{date_part}.txt"

    return send_file(buf, mimetype="text/plain", as_attachment=True, download_name=filename)


@app.route("/transcriptions/<int:trans_id>/download/docx")
@login_required
def download_docx(trans_id):
    """
    GET /transcriptions/<id>/download/docx
    Builds a .docx file in memory using python-docx and sends it as a download.
    python-docx lets us create Word documents without needing Microsoft Word installed.
    """
    conn = get_db()
    row = conn.execute(
        "SELECT text, created_at FROM transcriptions WHERE id = ? AND user_id = ?",
        (trans_id, session["user_id"])
    ).fetchone()
    conn.close()

    if not row:
        return jsonify({"error": "Not found."}), 404

    # Build the Word document
    doc = Document()

    # Title
    title = doc.add_heading("Transcription", level=1)
    title.runs[0].font.size = Pt(16)

    # Date subtitle
    date_part = (row["created_at"] or "")[:10]
    doc.add_paragraph(f"Date: {date_part}").runs[0].italic = True

    doc.add_paragraph("")  # blank line spacer

    # The transcription text — split into paragraphs on blank lines
    for para in row["text"].split("\n\n"):
        p = para.strip()
        if p:
            doc.add_paragraph(p)

    # Save the document to an in-memory buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    date_part = (row["created_at"] or "")[:10]
    filename = f"note-cloud-{date_part}.docx"

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ── Account stats & export ────────────────────────────────────────────────────

@app.route("/account/stats")
@login_required
def account_stats():
    """GET /account/stats — total transcription count and word count for the user."""
    conn = get_db()
    row = conn.execute(
        "SELECT COUNT(*) AS count, COALESCE(SUM(word_count), 0) AS words FROM transcriptions WHERE user_id = ?",
        (session["user_id"],)
    ).fetchone()
    user = conn.execute(
        "SELECT created_at FROM users WHERE id = ?", (session["user_id"],)
    ).fetchone()
    conn.close()
    return jsonify({
        "transcriptions": row["count"],
        "words": row["words"],
        "member_since": (user["created_at"] or "")[:10],
    })


@app.route("/account/export")
@login_required
def account_export():
    """GET /account/export — download all transcriptions as a JSON file."""
    import json as _json
    conn = get_db()
    rows = conn.execute(
        "SELECT id, title, text, word_count, created_at FROM transcriptions WHERE user_id = ? ORDER BY id DESC",
        (session["user_id"],)
    ).fetchall()
    conn.close()
    data = {"exported_at": datetime.utcnow().isoformat(), "transcriptions": [dict(r) for r in rows]}
    buf = io.BytesIO(_json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8"))
    buf.seek(0)
    return send_file(buf, mimetype="application/json", as_attachment=True,
                     download_name="note-cloud-export.json")


# ── Settings routes ───────────────────────────────────────────────────────────

@app.route("/profile/update", methods=["POST"])
@login_required
def profile_update():
    """
    POST /profile/update  { "first_name": "...", "last_name": "...", "email": "..." }
    Updates the user's display name and email address.
    """
    data = request.get_json(silent=True) or {}
    first_name = data.get("first_name", "").strip()
    last_name  = data.get("last_name",  "").strip()
    email      = data.get("email",      "").strip().lower()

    if not first_name or not last_name:
        return jsonify({"error": "First and last name are required."}), 400
    if not email or "@" not in email:
        return jsonify({"error": "A valid email address is required."}), 400

    conn = get_db()
    # Check if the new email is already taken by a different account
    existing = conn.execute(
        "SELECT id FROM users WHERE email = ? AND id != ?",
        (email, session["user_id"])
    ).fetchone()
    if existing:
        conn.close()
        return jsonify({"error": "That email is already in use by another account."}), 409

    conn.execute(
        "UPDATE users SET first_name = ?, last_name = ?, email = ? WHERE id = ?",
        (first_name, last_name, email, session["user_id"])
    )
    conn.commit()
    conn.close()

    session["user_email"] = email
    return jsonify({"ok": True, "first_name": first_name, "last_name": last_name, "email": email})


@app.route("/profile/password", methods=["POST"])
@login_required
def profile_password():
    """
    POST /profile/password  { "current": "...", "new_password": "...", "confirm": "..." }
    Changes the user's password after verifying their current one. Accounts
    created via "Continue with Google" have no password to verify — for those,
    this sets one for the first time instead of changing it.
    """
    data = request.get_json(silent=True) or {}
    current      = data.get("current",      "")
    new_password = data.get("new_password", "")
    confirm      = data.get("confirm",      "")

    if not new_password:
        return jsonify({"error": "New password is required."}), 400
    if len(new_password) < 6:
        return jsonify({"error": "New password must be at least 6 characters."}), 400
    if new_password != confirm:
        return jsonify({"error": "New passwords do not match."}), 400

    conn = get_db()
    user = conn.execute("SELECT password_hash, has_password FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    if not user:
        conn.close()
        return jsonify({"error": "Account not found."}), 404

    if user["has_password"]:
        if not current or not check_password_hash(user["password_hash"], current):
            conn.close()
            return jsonify({"error": "Current password is incorrect."}), 401

    new_hash = generate_password_hash(new_password, method="pbkdf2:sha256")
    conn.execute("UPDATE users SET password_hash = ?, has_password = 1 WHERE id = ?", (new_hash, session["user_id"]))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/account/delete", methods=["POST"])
@login_required
def account_delete():
    """
    POST /account/delete  { "password": "..." }
    Permanently deletes the account and all associated data after password
    confirmation. Accounts created via "Continue with Google" have no password
    to confirm — the authenticated session is treated as sufficient for those.
    """
    data = request.get_json(silent=True) or {}
    password = data.get("password", "")

    conn = get_db()
    user = conn.execute("SELECT password_hash, has_password FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    if not user:
        conn.close()
        return jsonify({"error": "Account not found."}), 404

    if user["has_password"]:
        if not password or not check_password_hash(user["password_hash"], password):
            conn.close()
            return jsonify({"error": "Incorrect password."}), 401

    uid = session["user_id"]
    # Delete all transcriptions, notebook memberships, and notebooks first
    trans_ids = [r["id"] for r in conn.execute("SELECT id FROM transcriptions WHERE user_id = ?", (uid,)).fetchall()]
    if trans_ids:
        placeholders = ",".join("?" * len(trans_ids))
        conn.execute(f"DELETE FROM notebook_transcriptions WHERE transcription_id IN ({placeholders})", trans_ids)
    conn.execute("DELETE FROM transcriptions WHERE user_id = ?", (uid,))
    conn.execute("DELETE FROM notebook_transcriptions WHERE notebook_id IN (SELECT id FROM notebooks WHERE user_id = ?)", (uid,))
    conn.execute("DELETE FROM notebooks WHERE user_id = ?", (uid,))
    conn.execute("DELETE FROM users WHERE id = ?", (uid,))
    conn.commit()
    conn.close()

    session.clear()
    return jsonify({"ok": True})


# ── Google Docs integration ───────────────────────────────────────────────────

def _google_flow():
    """Build a google_auth_oauthlib Flow from env config."""
    return Flow.from_client_config(
        {"web": {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [GOOGLE_REDIRECT_URI],
        }},
        scopes=GOOGLE_SCOPES,
        redirect_uri=GOOGLE_REDIRECT_URI,
    )


@app.route("/google/auth")
@login_required
def google_auth():
    """Redirect the user to Google's OAuth consent screen."""
    if not GOOGLE_LIBS_AVAILABLE or not GOOGLE_CLIENT_ID:
        return "Google integration not configured — add GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET to .env", 503
    flow = _google_flow()
    auth_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="consent",
    )
    session["google_oauth_state"] = state
    session["google_oauth_next"] = safe_next_path(request.args.get("next", "/"))
    return redirect(auth_url)


@app.route("/google/callback")
@login_required
def google_callback():
    """Handle the OAuth callback, store tokens, and redirect back to the app."""
    if not GOOGLE_LIBS_AVAILABLE or not GOOGLE_CLIENT_ID:
        return redirect("/?google_error=not_configured")

    state = session.pop("google_oauth_state", None)
    next_url = safe_next_path(session.pop("google_oauth_next", "/"))
    sep = "&" if "?" in next_url else "?"

    if not state or request.args.get("state") != state:
        return redirect(f"{next_url}{sep}google_error=invalid_state")

    if "error" in request.args:
        return redirect(f"{next_url}{sep}google_error=access_denied")

    flow = _google_flow()
    flow.fetch_token(authorization_response=request.url)
    creds = flow.credentials

    conn = get_db()
    conn.execute(
        "UPDATE users SET google_access_token=?, google_refresh_token=?, google_token_expiry=? WHERE id=?",
        (
            creds.token,
            creds.refresh_token,
            creds.expiry.isoformat() if creds.expiry else None,
            session["user_id"],
        ),
    )
    conn.commit()
    conn.close()
    return redirect(f"{next_url}{sep}google_connected=1")


@app.route("/google/status")
@login_required
def google_status():
    """Return whether the user has connected their Google account."""
    conn = get_db()
    user = conn.execute(
        "SELECT google_access_token FROM users WHERE id=?", (session["user_id"],)
    ).fetchone()
    conn.close()
    return jsonify({"connected": bool(user and user["google_access_token"])})


@app.route("/google/disconnect", methods=["POST"])
@login_required
def google_disconnect():
    """Remove stored Google tokens for the user."""
    conn = get_db()
    conn.execute(
        "UPDATE users SET google_access_token=NULL, google_refresh_token=NULL, google_token_expiry=NULL WHERE id=?",
        (session["user_id"],),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


@app.route("/transcriptions/<int:trans_id>/export/gdocs", methods=["POST"])
@login_required
def export_to_gdocs(trans_id):
    """Create a new Google Doc containing the transcription text."""
    if not GOOGLE_LIBS_AVAILABLE or not GOOGLE_CLIENT_ID:
        return jsonify({"error": "Google integration not configured."}), 503

    conn = get_db()
    row = conn.execute(
        "SELECT text, title, created_at FROM transcriptions WHERE id=? AND user_id=?",
        (trans_id, session["user_id"]),
    ).fetchone()
    user = conn.execute(
        "SELECT google_access_token, google_refresh_token, google_token_expiry FROM users WHERE id=?",
        (session["user_id"],),
    ).fetchone()

    if not row:
        conn.close()
        return jsonify({"error": "Not found."}), 404

    if not user or not user["google_access_token"]:
        conn.close()
        return jsonify({"error": "google_not_connected"}), 401

    # Build credentials object
    expiry = None
    if user["google_token_expiry"]:
        try:
            expiry = datetime.fromisoformat(user["google_token_expiry"])
        except Exception:
            pass

    creds = Credentials(
        token=user["google_access_token"],
        refresh_token=user["google_refresh_token"],
        token_uri="https://oauth2.googleapis.com/token",
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        scopes=GOOGLE_SCOPES,
        expiry=expiry,
    )

    # Refresh token if expired
    if creds.expired and creds.refresh_token:
        try:
            creds.refresh(GoogleRequest())
            conn.execute(
                "UPDATE users SET google_access_token=?, google_token_expiry=? WHERE id=?",
                (creds.token, creds.expiry.isoformat() if creds.expiry else None, session["user_id"]),
            )
            conn.commit()
        except Exception:
            conn.close()
            return jsonify({"error": "google_not_connected"}), 401

    conn.close()

    try:
        docs = google_build("docs", "v1", credentials=creds)
        doc_title = row["title"] or f"Transcription — {(row['created_at'] or '')[:10]}"
        doc = docs.documents().create(body={"title": doc_title}).execute()
        doc_id = doc["documentId"]

        text = (row["text"] or "").strip()
        if text:
            docs.documents().batchUpdate(
                documentId=doc_id,
                body={"requests": [{"insertText": {"location": {"index": 1}, "text": text}}]},
            ).execute()

        return jsonify({"ok": True, "url": f"https://docs.google.com/document/d/{doc_id}/edit"})
    except Exception as e:
        app.logger.error("gdocs export error: %s", e)
        return jsonify({"error": "Google Docs export failed — please try again."}), 500


# ── Notion integration ────────────────────────────────────────────────────────
# Uses a public Notion integration (OAuth), so each user connects their own
# workspace. Docs: https://developers.notion.com/docs/authorization

NOTION_API_BASE = "https://api.notion.com/v1"
NOTION_VERSION  = "2022-06-28"


def _notion_headers(token):
    return {
        "Authorization": f"Bearer {token}",
        "Notion-Version": NOTION_VERSION,
        "Content-Type": "application/json",
    }


@app.route("/notion/auth")
@login_required
def notion_auth():
    """Redirect the user to Notion's OAuth consent screen."""
    if not REQUESTS_AVAILABLE or not NOTION_CLIENT_ID:
        return "Notion integration not configured — add NOTION_CLIENT_ID and NOTION_CLIENT_SECRET to .env", 503

    state = secrets.token_urlsafe(24)
    session["notion_oauth_state"] = state

    session["notion_oauth_next"] = safe_next_path(request.args.get("next", "/"))

    from urllib.parse import urlencode
    params = {
        "client_id": NOTION_CLIENT_ID,
        "response_type": "code",
        "owner": "user",
        "redirect_uri": NOTION_REDIRECT_URI,
        "state": state,
    }
    return redirect(f"{NOTION_API_BASE}/oauth/authorize?{urlencode(params)}")


@app.route("/notion/callback")
@login_required
def notion_callback():
    """Handle the OAuth callback, store the workspace token, and redirect back."""
    next_url = safe_next_path(session.pop("notion_oauth_next", "/"))
    sep = "&" if "?" in next_url else "?"

    if not REQUESTS_AVAILABLE or not NOTION_CLIENT_ID:
        return redirect(f"{next_url}{sep}notion_error=not_configured")

    state = session.pop("notion_oauth_state", None)
    if not state or request.args.get("state") != state:
        return redirect(f"{next_url}{sep}notion_error=invalid_state")

    if "error" in request.args:
        return redirect(f"{next_url}{sep}notion_error=access_denied")

    code = request.args.get("code")
    if not code:
        return redirect(f"{next_url}{sep}notion_error=access_denied")

    try:
        resp = http_requests.post(
            f"{NOTION_API_BASE}/oauth/token",
            auth=(NOTION_CLIENT_ID, NOTION_CLIENT_SECRET),
            json={"grant_type": "authorization_code", "code": code, "redirect_uri": NOTION_REDIRECT_URI},
            timeout=10,
        )
        resp.raise_for_status()
        payload = resp.json()
    except Exception as e:
        app.logger.error("notion token exchange error: %s", e)
        return redirect(f"{next_url}{sep}notion_error=access_denied")

    conn = get_db()
    conn.execute(
        "UPDATE users SET notion_access_token=?, notion_workspace_name=?, notion_bot_id=? WHERE id=?",
        (
            payload.get("access_token"),
            payload.get("workspace_name"),
            payload.get("bot_id"),
            session["user_id"],
        ),
    )
    conn.commit()
    conn.close()
    return redirect(f"{next_url}{sep}notion_connected=1")


@app.route("/notion/status")
@login_required
def notion_status():
    """Return whether the user has connected Notion, and to which workspace."""
    conn = get_db()
    user = conn.execute(
        "SELECT notion_access_token, notion_workspace_name FROM users WHERE id=?", (session["user_id"],)
    ).fetchone()
    conn.close()
    return jsonify({
        "connected": bool(user and user["notion_access_token"]),
        "workspace": user["notion_workspace_name"] if user else None,
    })


@app.route("/notion/disconnect", methods=["POST"])
@login_required
def notion_disconnect():
    """Remove the stored Notion token for the user."""
    conn = get_db()
    conn.execute(
        "UPDATE users SET notion_access_token=NULL, notion_workspace_name=NULL, notion_bot_id=NULL WHERE id=?",
        (session["user_id"],),
    )
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


def _notion_text_blocks(text):
    """Split text into Notion paragraph blocks, chunked under the 2000-char rich-text limit."""
    blocks = []
    for para in (text or "").split("\n\n"):
        para = para.strip("\n")
        if not para:
            continue
        for i in range(0, len(para), 1900):
            chunk = para[i:i + 1900]
            blocks.append({
                "object": "block",
                "type": "paragraph",
                "paragraph": {"rich_text": [{"type": "text", "text": {"content": chunk}}]},
            })
    return blocks or [{"object": "block", "type": "paragraph", "paragraph": {"rich_text": []}}]


@app.route("/transcriptions/<int:trans_id>/export/notion", methods=["POST"])
@login_required
def export_to_notion(trans_id):
    """Create a new Notion page containing the transcription text."""
    if not REQUESTS_AVAILABLE or not NOTION_CLIENT_ID:
        return jsonify({"error": "Notion integration not configured."}), 503

    conn = get_db()
    row = conn.execute(
        "SELECT text, title, created_at FROM transcriptions WHERE id=? AND user_id=?",
        (trans_id, session["user_id"]),
    ).fetchone()
    user = conn.execute(
        "SELECT notion_access_token FROM users WHERE id=?", (session["user_id"],)
    ).fetchone()
    conn.close()

    if not row:
        return jsonify({"error": "Not found."}), 404
    if not user or not user["notion_access_token"]:
        return jsonify({"error": "notion_not_connected"}), 401

    token = user["notion_access_token"]
    headers = _notion_headers(token)

    try:
        # Notion has no "workspace root" you can create pages under directly —
        # a new page must live inside a page or database the user has already
        # shared with the integration during the connect flow. Find one.
        search_resp = http_requests.post(
            f"{NOTION_API_BASE}/search",
            headers=headers,
            json={"filter": {"value": "page", "property": "object"}, "page_size": 1},
            timeout=10,
        )
        search_resp.raise_for_status()
        results = search_resp.json().get("results", [])
        if not results:
            return jsonify({
                "error": "notion_no_pages",
                "message": "Connected, but no pages are shared with Note-Cloud yet. In Notion, open a page, "
                           "click ⋯ → Connections, and add Note-Cloud — then try again.",
            }), 422
        parent_id = results[0]["id"]

        doc_title = row["title"] or f"Transcription — {(row['created_at'] or '')[:10]}"
        create_resp = http_requests.post(
            f"{NOTION_API_BASE}/pages",
            headers=headers,
            json={
                "parent": {"page_id": parent_id},
                "properties": {"title": {"title": [{"type": "text", "text": {"content": doc_title}}]}},
                "children": _notion_text_blocks(row["text"]),
            },
            timeout=10,
        )
        create_resp.raise_for_status()
        page = create_resp.json()
        url = page.get("url") or f"https://notion.so/{page['id'].replace('-', '')}"
        return jsonify({"ok": True, "url": url})
    except Exception as e:
        app.logger.error("notion export error: %s", e)
        return jsonify({"error": "Notion export failed — please try again."}), 500


# ── Entry point ────────────────────────────────────────────────────────────────

# ── Google Login (OAuth for authentication) ───────────────────────────────────

@app.route("/auth/google")
def google_login():
    """Redirect to Google's consent screen for sign-in."""
    if not GOOGLE_LIBS_AVAILABLE or not GOOGLE_CLIENT_ID:
        return redirect("/login?error=google_not_configured")
    flow = Flow.from_client_config(
        {"web": {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [GOOGLE_LOGIN_REDIRECT_URI],
        }},
        scopes=GOOGLE_LOGIN_SCOPES,
        redirect_uri=GOOGLE_LOGIN_REDIRECT_URI,
    )
    auth_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="select_account",
    )
    session["google_login_state"] = state
    session["google_login_verifier"] = flow.code_verifier
    return redirect(auth_url)


@app.route("/auth/google/callback")
def google_login_callback():
    """Exchange code for tokens, fetch profile, create or log in user."""
    if not GOOGLE_LIBS_AVAILABLE or not GOOGLE_CLIENT_ID:
        return redirect("/login?error=google_not_configured")

    state = session.pop("google_login_state", None)
    if not state or request.args.get("state") != state:
        return redirect("/login?error=invalid_state")

    if "error" in request.args:
        return redirect("/login?error=access_denied")

    flow = Flow.from_client_config(
        {"web": {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [GOOGLE_LOGIN_REDIRECT_URI],
        }},
        scopes=GOOGLE_LOGIN_SCOPES,
        redirect_uri=GOOGLE_LOGIN_REDIRECT_URI,
        state=state,
    )
    flow.code_verifier = session.pop("google_login_verifier", None)
    flow.fetch_token(authorization_response=request.url)
    creds = flow.credentials

    # Fetch user profile from Google
    import urllib.request, json as _json
    req = urllib.request.Request(
        "https://www.googleapis.com/oauth2/v3/userinfo",
        headers={"Authorization": f"Bearer {creds.token}"}
    )
    with urllib.request.urlopen(req) as resp:
        profile = _json.loads(resp.read().decode())

    email      = profile.get("email", "").lower().strip()
    first_name = profile.get("given_name", "")
    last_name  = profile.get("family_name", "")
    avatar_url = profile.get("picture", "")

    if not email:
        return redirect("/login?error=no_email")

    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()

    if user:
        # Existing account — just log in
        session["user_id"]    = user["id"]
        session["user_email"] = user["email"]
        conn.close()
    else:
        # New account via Google — create with random password
        pw_hash      = generate_password_hash(secrets.token_hex(32), method="pbkdf2:sha256")
        created_at   = datetime.utcnow().isoformat()
        ref_code     = secrets.token_urlsafe(6).upper()[:8]
        try:
            cursor = conn.execute(
                """INSERT INTO users
                   (email, password_hash, first_name, last_name, created_at, referral_code, tier, has_password)
                   VALUES (?, ?, ?, ?, ?, ?, 'free', 0)""",
                (email, pw_hash, first_name or "User", last_name or "", created_at, ref_code),
            )
            conn.commit()
            new_id = cursor.lastrowid
        except Exception:
            conn.close()
            return redirect("/login?error=account_error")

        session["user_id"]    = new_id
        session["user_email"] = email
        conn.close()

    return redirect("/")


# ── Apple Sign-In (Sign in with Apple) ────────────────────────────────────────
# Docs: https://developer.apple.com/documentation/sign_in_with_apple

def _apple_configured():
    return bool(APPLE_LIBS_AVAILABLE and APPLE_CLIENT_ID and APPLE_TEAM_ID and APPLE_KEY_ID and APPLE_PRIVATE_KEY)


@app.route("/auth/apple")
def apple_login():
    """Redirect to Apple's Sign in with Apple consent screen."""
    if not _apple_configured():
        return redirect("/login?error=apple_not_configured")

    state = secrets.token_urlsafe(24)
    session["apple_oauth_state"] = state

    from urllib.parse import urlencode
    params = {
        "client_id": APPLE_CLIENT_ID,
        "redirect_uri": APPLE_REDIRECT_URI,
        "response_type": "code id_token",
        "response_mode": "form_post",
        "scope": "name email",
        "state": state,
    }
    return redirect(f"https://appleid.apple.com/auth/authorize?{urlencode(params)}")


@app.route("/auth/apple/callback", methods=["POST"])
@csrf.exempt  # Apple posts here directly — no CSRF token to send. The signed
              # id_token + our own state check are what verify authenticity.
def apple_login_callback():
    """
    Apple posts here (not a GET redirect) with the signed id_token, an
    authorization code, and — only on the very first authorization ever — a
    `user` field with the name Apple only shares once.
    """
    if not _apple_configured():
        return redirect("/login?error=apple_not_configured")

    state = session.pop("apple_oauth_state", None)
    if not state or request.form.get("state") != state:
        return redirect("/login?error=invalid_state")

    if "error" in request.form:
        return redirect("/login?error=access_denied")

    id_token = request.form.get("id_token")
    if not id_token:
        return redirect("/login?error=access_denied")

    try:
        jwks_client = pyjwt.PyJWKClient("https://appleid.apple.com/auth/keys")
        signing_key = jwks_client.get_signing_key_from_jwt(id_token)
        claims = pyjwt.decode(
            id_token,
            signing_key.key,
            algorithms=["RS256"],
            audience=APPLE_CLIENT_ID,
            issuer="https://appleid.apple.com",
        )
    except Exception as e:
        app.logger.error("apple id_token verification failed: %s", e)
        return redirect("/login?error=access_denied")

    apple_sub = claims.get("sub")
    email     = (claims.get("email") or "").lower().strip()
    if not apple_sub:
        return redirect("/login?error=access_denied")

    # Apple only sends the user's name once, on their very first authorization.
    first_name, last_name = "", ""
    user_json = request.form.get("user")
    if user_json:
        try:
            name = json.loads(user_json).get("name", {})
            first_name = name.get("firstName", "")
            last_name  = name.get("lastName", "")
        except Exception:
            pass

    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE apple_sub = ?", (apple_sub,)).fetchone()

    if not user and email:
        # First Apple login for an email that already has a password account —
        # link them instead of creating a duplicate.
        user = conn.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
        if user:
            conn.execute("UPDATE users SET apple_sub = ? WHERE id = ?", (apple_sub, user["id"]))
            conn.commit()

    if user:
        session["user_id"]    = user["id"]
        session["user_email"] = user["email"]
        conn.close()
    else:
        if not email:
            conn.close()
            return redirect("/login?error=no_email")
        pw_hash    = generate_password_hash(secrets.token_hex(32), method="pbkdf2:sha256")
        created_at = datetime.utcnow().isoformat()
        ref_code   = secrets.token_urlsafe(6).upper()[:8]
        try:
            cursor = conn.execute(
                """INSERT INTO users
                   (email, password_hash, first_name, last_name, created_at, referral_code, tier, has_password, apple_sub)
                   VALUES (?, ?, ?, ?, ?, ?, 'free', 0, ?)""",
                (email, pw_hash, first_name or "User", last_name or "", created_at, ref_code, apple_sub),
            )
            conn.commit()
            new_id = cursor.lastrowid
        except Exception:
            conn.close()
            return redirect("/login?error=account_error")

        session["user_id"]    = new_id
        session["user_email"] = email
        conn.close()

    return redirect("/")


if __name__ == "__main__":
    # Debug mode exposes the Werkzeug interactive debugger, which allows
    # arbitrary code execution — it must NEVER be on in production. Default to
    # on for local dev, but force it off when FLASK_ENV=production.
    debug = os.getenv("FLASK_ENV") != "production"
    app.run(debug=debug, port=5000)
